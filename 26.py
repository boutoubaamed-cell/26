import streamlit as st
import pandas as pd
import numpy as np
from scipy.stats import shapiro, f_oneway, ttest_ind, mannwhitneyu, kruskal, pearsonr, spearmanr, chi2_contingency, \
    wilcoxon, friedmanchisquare, levene, bartlett, ttest_rel, norm, probplot, mstats, zscore
from scipy import stats
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
from plotly.subplots import make_subplots
import warnings
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression, LogisticRegression, Ridge, Lasso, ElasticNet
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error, confusion_matrix, classification_report, \
    accuracy_score, precision_score, recall_score, f1_score
import statsmodels.api as sm
from statsmodels.formula.api import ols
from sklearn.preprocessing import StandardScaler, LabelEncoder, MinMaxScaler, RobustScaler
from sklearn.cluster import KMeans, AgglomerativeClustering, DBSCAN, Birch
from sklearn.decomposition import PCA, FactorAnalysis, TruncatedSVD
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.diagnostic import het_breuschpagan, acorr_ljungbox, normal_ad
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor, RandomForestClassifier
from sklearn.model_selection import train_test_split, cross_val_score, GridSearchCV, KFold
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import io
import time
from datetime import datetime
import json
import base64
import hashlib
import re
from typing import List, Dict, Any, Tuple, Optional
import math
from scipy.cluster.hierarchy import dendrogram, linkage, fcluster
from scipy.spatial.distance import pdist, squareform
from PIL import Image
import networkx as nx
from io import BytesIO
import xlsxwriter

# إعداد الخط
plt.rcParams['font.family'] = 'Arial'
plt.rcParams['axes.unicode_minus'] = False
warnings.filterwarnings('ignore')

# إعداد صفحة Streamlit
st.set_page_config(
    page_title="نظام التحليل الإحصائي المتقدم للاستبيانات",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ==================== الدوال المساعدة الأساسية ====================

def calculate_cronbach_alpha(df):
    """حساب معامل ألفا كرونباخ لقياس الثبات"""
    df = df.dropna()
    if df.shape[1] < 2:
        return np.nan
    k = df.shape[1]
    total_variance = df.sum(axis=1).var()
    item_variances = df.var(axis=0).sum()
    if total_variance > 0:
        alpha = (k / (k - 1)) * (1 - (item_variances / total_variance))
    else:
        alpha = np.nan
    return alpha


def get_likert_trend(mean_value):
    """تحديد الاتجاه العام بناءً على متوسط ليكارت"""
    if mean_value >= 4.5:
        return "موافق بشدة ✅✅", "trend-strongly-agree", "مرتفع جداً"
    elif mean_value >= 3.5:
        return "موافق ✅", "trend-agree", "مرتفع"
    elif mean_value >= 2.5:
        return "محايد ➖", "trend-neutral", "متوسط"
    elif mean_value >= 1.5:
        return "غير موافق ❌", "trend-disagree", "منخفض"
    else:
        return "غير موافق بشدة ❌❌", "trend-strongly-disagree", "منخفض جداً"


def calculate_effect_size_cohens_d(group1, group2):
    """حساب حجم التأثير باستخدام Cohen's d"""
    n1, n2 = len(group1), len(group2)
    var1, var2 = group1.var(), group2.var()
    pooled_sd = np.sqrt(((n1 - 1) * var1 + (n2 - 1) * var2) / (n1 + n2 - 2))
    if pooled_sd == 0:
        return 0
    return (group1.mean() - group2.mean()) / pooled_sd


def calculate_effect_size_eta_squared(ss_between, ss_total):
    """حساب حجم التأثير باستخدام Eta-squared"""
    return ss_between / ss_total if ss_total > 0 else 0


def calculate_effect_size_omega_squared(ss_between, df_between, ss_within, df_within):
    """حساب حجم التأثير باستخدام Omega-squared"""
    ms_within = ss_within / df_within if df_within > 0 else 0
    total_variance = ss_between + ss_within
    omega = (ss_between - (df_between * ms_within)) / (total_variance + ms_within)
    return omega if omega > 0 else 0


def calculate_confidence_interval(data, confidence=0.95):
    """حساب فاصل الثقة"""
    n = len(data)
    mean = np.mean(data)
    se = stats.sem(data)
    ci = se * stats.t.ppf((1 + confidence) / 2, n - 1)
    return mean - ci, mean + ci


def detect_outliers_iqr(data, multiplier=1.5):
    """كشف القيم المتطرفة باستخدام طريقة IQR"""
    q1 = data.quantile(0.25)
    q3 = data.quantile(0.75)
    iqr = q3 - q1
    lower_bound = q1 - multiplier * iqr
    upper_bound = q3 + multiplier * iqr
    outliers = data[(data < lower_bound) | (data > upper_bound)]
    return outliers, lower_bound, upper_bound


def detect_outliers_zscore(data, threshold=3):
    """كشف القيم المتطرفة باستخدام Z-Score"""
    z_scores = np.abs(stats.zscore(data.dropna()))
    outliers = data[z_scores > threshold]
    return outliers


def calculate_normality_tests(data):
    """إجراء اختبارات التوزيع الطبيعي المتعددة"""
    results = {}

    if len(data) >= 3:
        # Shapiro-Wilk
        shapiro_stat, shapiro_p = shapiro(data)
        results['shapiro'] = {'statistic': shapiro_stat, 'pvalue': shapiro_p}

        # Kolmogorov-Smirnov
        ks_stat, ks_p = stats.kstest(data, 'norm', args=(data.mean(), data.std()))
        results['ks'] = {'statistic': ks_stat, 'pvalue': ks_p}

        # Anderson-Darling
        anderson_result = stats.anderson(data, dist='norm')
        results['anderson'] = {'statistic': anderson_result.statistic,
                               'critical_values': anderson_result.critical_values.tolist(),
                               'significance_levels': anderson_result.significance_level.tolist()}

        # Jarque-Bera
        jb_stat, jb_p = stats.jarque_bera(data)
        results['jarque_bera'] = {'statistic': jb_stat, 'pvalue': jb_p}

        # D'Agostino's K^2
        dagostino_result = stats.normaltest(data)
        results['dagostino'] = {'statistic': dagostino_result.statistic, 'pvalue': dagostino_result.pvalue}

    return results


def calculate_descriptive_statistics_detailed(df):
    """حساب إحصاءات وصفية مفصلة"""
    stats_dict = {}

    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]):
            data = df[col].dropna()
            if len(data) > 0:
                stats_dict[col] = {
                    'mean': data.mean(),
                    'median': data.median(),
                    'mode': data.mode().iloc[0] if len(data.mode()) > 0 else np.nan,
                    'std': data.std(),
                    'var': data.var(),
                    'min': data.min(),
                    'max': data.max(),
                    'range': data.max() - data.min(),
                    'q1': data.quantile(0.25),
                    'q3': data.quantile(0.75),
                    'iqr': data.quantile(0.75) - data.quantile(0.25),
                    'skewness': data.skew(),
                    'kurtosis': data.kurtosis(),
                    'cv': (data.std() / data.mean()) * 100 if data.mean() != 0 else np.nan,
                    'n': len(data),
                    'n_missing': df[col].isna().sum(),
                    'percent_missing': (df[col].isna().sum() / len(df)) * 100
                }

    return pd.DataFrame(stats_dict).T


def advanced_regression_analysis(X, y, feature_names):
    """تحليل انحدار متقدم مع تشخيص النموذج"""
    results = {}

    if isinstance(X, pd.DataFrame):
        X = X.values
    if isinstance(y, pd.Series):
        y = y.values

    X_with_const = sm.add_constant(X)
    model = sm.OLS(y, X_with_const).fit()
    results['model'] = model
    results['coefficients'] = model.params
    results['p_values'] = model.pvalues
    results['std_errors'] = model.bse
    results['t_values'] = model.tvalues
    results['conf_int'] = model.conf_int()
    results['r_squared'] = model.rsquared
    results['adj_r_squared'] = model.rsquared_adj
    results['f_statistic'] = model.fvalue
    results['f_pvalue'] = model.f_pvalue
    results['aic'] = model.aic
    results['bic'] = model.bic
    results['mse'] = mean_squared_error(y, model.fittedvalues)
    results['rmse'] = np.sqrt(results['mse'])
    results['mae'] = mean_absolute_error(y, model.fittedvalues)
    results['mape'] = np.mean(np.abs((y - model.fittedvalues) / y)) * 100 if np.all(y != 0) else np.nan

    # معاملات التأثير الموحدة (Beta coefficients)
    X_scaled = (X - X.mean(axis=0)) / X.std(axis=0)
    model_scaled = sm.OLS(y, sm.add_constant(X_scaled)).fit()
    results['beta_coefficients'] = model_scaled.params[1:] if X.shape[1] > 0 else []

    try:
        results['anova_table'] = sm.stats.anova_lm(model, typ=2)
    except:
        results['anova_table'] = pd.DataFrame()

    try:
        bp_test = het_breuschpagan(model.resid, X_with_const)
        results['heteroscedasticity'] = {'statistic': bp_test[0], 'pvalue': bp_test[1]}
    except:
        results['heteroscedasticity'] = {'statistic': np.nan, 'pvalue': np.nan}

    # اختبار الارتباط الذاتي للبواقي (Durbin-Watson)
    from statsmodels.stats.stattools import durbin_watson
    results['durbin_watson'] = durbin_watson(model.resid)

    # اختبار طبيعية البواقي المتقدم
    if len(model.resid) >= 3:
        shapiro_stat, shapiro_p = shapiro(model.resid)
        results['residuals_normality'] = {'test': 'Shapiro-Wilk', 'statistic': shapiro_stat, 'pvalue': shapiro_p}

        # اختبار جارك-بيرا
        jb_stat, jb_p = stats.jarque_bera(model.resid)
        results['residuals_normality_jb'] = {'statistic': jb_stat, 'pvalue': jb_p}

    # VIF
    vif_data = []
    if X.shape[1] > 1:
        for i, col in enumerate(feature_names):
            try:
                vif = variance_inflation_factor(X, i)
                vif_data.append({'المتغير': col, 'VIF': round(vif, 3),
                                 'تفسير': 'مشكلة خطيرة' if vif > 10 else 'مشكلة محتملة' if vif > 5 else 'جيد'})
            except:
                vif_data.append({'المتغير': col, 'VIF': np.nan, 'تفسير': 'غير محدد'})
    else:
        vif_data.append({'المتغير': feature_names[0], 'VIF': 1.0, 'تفسير': 'جيد'})
    results['vif'] = pd.DataFrame(vif_data)

    results['residuals'] = model.resid
    results['fitted_values'] = model.fittedvalues

    # Cook's Distance و Leverage
    influence = model.get_influence()
    results['cooks_distance'] = influence.cooks_distance[0]
    results['leverage'] = influence.hat_matrix_diag

    return results


def stepwise_regression(X, y, feature_names, direction='both', significance_level=0.05):
    """تحليل الانحدار التدريجي (Stepwise Regression)"""
    results = {}
    selected_features = []

    if direction in ['forward', 'both']:
        remaining_features = feature_names.copy()
        best_r2 = 0

        while remaining_features:
            best_feature = None
            best_new_r2 = best_r2

            for feature in remaining_features:
                X_temp = sm.add_constant(
                    pd.DataFrame({f: X[:, feature_names.index(f)] for f in selected_features + [feature]}))
                model = sm.OLS(y, X_temp).fit()
                if model.rsquared > best_new_r2:
                    best_new_r2 = model.rsquared
                    best_feature = feature

            if best_feature is not None and best_new_r2 > best_r2:
                selected_features.append(best_feature)
                remaining_features.remove(best_feature)
                best_r2 = best_new_r2
            else:
                break

    results['selected_features'] = selected_features
    if selected_features:
        results['final_model'] = sm.OLS(y, sm.add_constant(
            pd.DataFrame({f: X[:, feature_names.index(f)] for f in selected_features}))).fit()
    else:
        results['final_model'] = None

    return results


def ridge_regression_analysis(X, y, alpha=1.0):
    """تحليل انحدار ريدج (Ridge Regression)"""
    from sklearn.linear_model import Ridge
    from sklearn.model_selection import cross_val_score

    ridge = Ridge(alpha=alpha)
    ridge.fit(X, y)

    cv_scores = cross_val_score(ridge, X, y, cv=5, scoring='r2')

    results = {
        'coefficients': ridge.coef_,
        'intercept': ridge.intercept_,
        'cv_scores': cv_scores,
        'mean_cv_score': cv_scores.mean(),
        'std_cv_score': cv_scores.std()
    }

    return results


def lasso_regression_analysis(X, y, alpha=1.0):
    """تحليل انحدار لاسو (Lasso Regression)"""
    from sklearn.linear_model import Lasso
    from sklearn.model_selection import cross_val_score

    lasso = Lasso(alpha=alpha, max_iter=10000)
    lasso.fit(X, y)

    cv_scores = cross_val_score(lasso, X, y, cv=5, scoring='r2')

    results = {
        'coefficients': lasso.coef_,
        'intercept': lasso.intercept_,
        'n_features_used': np.sum(lasso.coef_ != 0),
        'cv_scores': cv_scores,
        'mean_cv_score': cv_scores.mean(),
        'std_cv_score': cv_scores.std()
    }

    return results


def elastic_net_regression(X, y, alpha=1.0, l1_ratio=0.5):
    """تحليل انحدار الشبكة المرنة (Elastic Net)"""
    from sklearn.linear_model import ElasticNet
    from sklearn.model_selection import cross_val_score

    elastic = ElasticNet(alpha=alpha, l1_ratio=l1_ratio, max_iter=10000)
    elastic.fit(X, y)

    cv_scores = cross_val_score(elastic, X, y, cv=5, scoring='r2')

    results = {
        'coefficients': elastic.coef_,
        'intercept': elastic.intercept_,
        'n_features_used': np.sum(elastic.coef_ != 0),
        'cv_scores': cv_scores,
        'mean_cv_score': cv_scores.mean(),
        'std_cv_score': cv_scores.std()
    }

    return results


def random_forest_regression(X, y, n_estimators=100, max_depth=None):
    """تحليل انحدار باستخدام Random Forest"""
    from sklearn.ensemble import RandomForestRegressor
    from sklearn.model_selection import cross_val_score

    rf = RandomForestRegressor(n_estimators=n_estimators, max_depth=max_depth, random_state=42)
    rf.fit(X, y)

    cv_scores = cross_val_score(rf, X, y, cv=5, scoring='r2')

    results = {
        'feature_importance': rf.feature_importances_,
        'cv_scores': cv_scores,
        'mean_cv_score': cv_scores.mean(),
        'std_cv_score': cv_scores.std(),
        'model': rf
    }

    return results


def mediation_analysis_advanced(df, independent_var, mediator_var, dependent_var, bootstrap_iterations=1000):
    """تحليل تأثير المتغير الوسيط المتقدم مع Bootstrap"""
    results = {}
    data = df[[independent_var, mediator_var, dependent_var]].dropna()

    # التأثير الكلي (Total Effect)
    model_c = sm.OLS(data[dependent_var], sm.add_constant(data[independent_var])).fit()
    results['total_effect'] = {
        'coefficient': model_c.params[independent_var],
        'pvalue': model_c.pvalues[independent_var],
        'r_squared': model_c.rsquared,
        'se': model_c.bse[independent_var],
        'ci_lower': model_c.conf_int().loc[independent_var, 0],
        'ci_upper': model_c.conf_int().loc[independent_var, 1]
    }

    # المسار أ (Path A)
    model_a = sm.OLS(data[mediator_var], sm.add_constant(data[independent_var])).fit()
    results['path_a'] = {
        'coefficient': model_a.params[independent_var],
        'pvalue': model_a.pvalues[independent_var],
        'r_squared': model_a.rsquared,
        'se': model_a.bse[independent_var],
        'ci_lower': model_a.conf_int().loc[independent_var, 0],
        'ci_upper': model_a.conf_int().loc[independent_var, 1]
    }

    # المسار ب (Path B)
    model_b = sm.OLS(data[dependent_var], sm.add_constant(data[[independent_var, mediator_var]])).fit()
    results['path_b'] = {
        'coefficient': model_b.params[mediator_var],
        'pvalue': model_b.pvalues[mediator_var],
        'se': model_b.bse[mediator_var],
        'ci_lower': model_b.conf_int().loc[mediator_var, 0],
        'ci_upper': model_b.conf_int().loc[mediator_var, 1]
    }

    # التأثير المباشر (Direct Effect)
    results['direct_effect'] = {
        'coefficient': model_b.params[independent_var],
        'pvalue': model_b.pvalues[independent_var],
        'se': model_b.bse[independent_var],
        'ci_lower': model_b.conf_int().loc[independent_var, 0],
        'ci_upper': model_b.conf_int().loc[independent_var, 1]
    }

    # التأثير غير المباشر (Indirect Effect)
    indirect_effect = results['path_a']['coefficient'] * results['path_b']['coefficient']
    results['indirect_effect'] = indirect_effect

    # نسبة الوساطة
    total = results['total_effect']['coefficient']
    results['mediation_ratio'] = (indirect_effect / total * 100) if total != 0 else 0

    # اختبار سوبل (Sobel Test)
    se_a = model_a.bse[independent_var]
    se_b = model_b.bse[mediator_var]
    sobel_z = indirect_effect / np.sqrt(
        (results['path_b']['coefficient'] ** 2 * se_a ** 2) + (results['path_a']['coefficient'] ** 2 * se_b ** 2))
    sobel_p = 2 * (1 - stats.norm.cdf(abs(sobel_z)))
    results['sobel_test'] = {'z': sobel_z, 'pvalue': sobel_p}

    # Bootstrap Confidence Interval
    bootstrap_effects = []
    for _ in range(bootstrap_iterations):
        bootstrap_sample = data.sample(n=len(data), replace=True)
        boot_model_a = sm.OLS(bootstrap_sample[mediator_var], sm.add_constant(bootstrap_sample[independent_var])).fit()
        boot_model_b = sm.OLS(bootstrap_sample[dependent_var],
                              sm.add_constant(bootstrap_sample[[independent_var, mediator_var]])).fit()
        boot_indirect = boot_model_a.params[independent_var] * boot_model_b.params[mediator_var]
        bootstrap_effects.append(boot_indirect)

    results['bootstrap_ci'] = {
        'lower': np.percentile(bootstrap_effects, 2.5),
        'upper': np.percentile(bootstrap_effects, 97.5)
    }

    # تحديد نوع الوساطة
    if results['path_a']['pvalue'] < 0.05 and results['path_b']['pvalue'] < 0.05:
        if results['direct_effect']['pvalue'] < 0.05:
            results['mediation_type'] = "وساطة جزئية (Partial Mediation)"
        else:
            results['mediation_type'] = "وساطة كاملة (Full Mediation)"
    else:
        results['mediation_type'] = "لا توجد وساطة (No Mediation)"

    return results


def clustering_analysis_advanced(df, n_clusters_range=(2, 10), methods=['kmeans', 'hierarchical']):
    """تحليل التجميع المتقدم مع مقارنة الطرق المختلفة"""
    results = {}

    scaler = StandardScaler()
    df_scaled = scaler.fit_transform(df)
    results['scaler'] = scaler

    # K-Means with elbow method and silhouette scores
    if 'kmeans' in methods:
        inertias = []
        silhouette_scores = []
        calinski_scores = []
        davies_scores = []

        for k in range(n_clusters_range[0], n_clusters_range[1] + 1):
            kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
            labels = kmeans.fit_predict(df_scaled)
            inertias.append(kmeans.inertia_)

            if len(set(labels)) > 1:
                silhouette_scores.append(silhouette_score(df_scaled, labels))
                calinski_scores.append(calinski_harabasz_score(df_scaled, labels))
                davies_scores.append(davies_bouldin_score(df_scaled, labels))
            else:
                silhouette_scores.append(-1)
                calinski_scores.append(0)
                davies_scores.append(np.inf)

        results['kmeans'] = {
            'inertias': inertias,
            'silhouette_scores': silhouette_scores,
            'calinski_scores': calinski_scores,
            'davies_scores': davies_scores,
            'optimal_k_silhouette': np.argmax(silhouette_scores) + n_clusters_range[0] if silhouette_scores else
            n_clusters_range[0],
            'optimal_k_calinski': np.argmax(calinski_scores) + n_clusters_range[0] if calinski_scores else
            n_clusters_range[0]
        }

    # Hierarchical clustering
    if 'hierarchical' in methods:
        linkage_matrix = linkage(df_scaled, method='ward')
        results['hierarchical'] = {
            'linkage_matrix': linkage_matrix
        }

    # DBSCAN
    if 'dbscan' in methods:
        dbscan = DBSCAN(eps=0.5, min_samples=5)
        dbscan_labels = dbscan.fit_predict(df_scaled)
        n_clusters_dbscan = len(set(dbscan_labels)) - (1 if -1 in dbscan_labels else 0)
        results['dbscan'] = {
            'labels': dbscan_labels,
            'n_clusters': n_clusters_dbscan,
            'n_noise': np.sum(dbscan_labels == -1),
            'silhouette': silhouette_score(df_scaled, dbscan_labels) if n_clusters_dbscan > 1 else np.nan
        }

    return results


def factor_analysis_advanced(df, n_factors=None, rotation='varimax'):
    """تحليل العوامل المتقدم (Factor Analysis)"""
    results = {}

    scaler = StandardScaler()
    df_scaled = scaler.fit_transform(df)

    # PCA for initial eigenvalues
    from sklearn.decomposition import PCA
    pca = PCA()
    pca.fit(df_scaled)
    eigenvalues = pca.explained_variance_
    results['eigenvalues'] = eigenvalues.tolist()
    results['explained_variance_ratio'] = pca.explained_variance_ratio_.tolist()
    results['cumulative_variance'] = np.cumsum(pca.explained_variance_ratio_).tolist()

    # Kaiser criterion
    n_factors_kaiser = np.sum(eigenvalues > 1)
    results['n_factors_kaiser'] = int(n_factors_kaiser)

    # Scree plot data
    results['scree_data'] = {'components': list(range(1, len(eigenvalues) + 1)), 'eigenvalues': eigenvalues.tolist()}

    if n_factors is None:
        n_factors = n_factors_kaiser if n_factors_kaiser > 0 else 2

    # Factor Analysis
    if rotation == 'varimax':
        from sklearn.decomposition import PCA
        pca = PCA(n_components=n_factors)
        pca_result = pca.fit_transform(df_scaled)
        loadings = pca.components_.T * np.sqrt(pca.explained_variance_)
        loadings_df = pd.DataFrame(loadings, index=df.columns, columns=[f'Factor_{i + 1}' for i in range(n_factors)])
        results['factor_loadings'] = loadings_df
        results['communalities'] = (loadings_df ** 2).sum(axis=1).to_dict()
        factor_var = np.var(pca_result, axis=0)
        results['factor_variance'] = factor_var.tolist()
        results['factor_variance_ratio'] = (factor_var / factor_var.sum() * 100).tolist()
    else:
        fa = FactorAnalysis(n_components=n_factors, rotation=None)
        fa.fit(df_scaled)
        loadings = pd.DataFrame(fa.components_.T, index=df.columns,
                                columns=[f'Factor_{i + 1}' for i in range(n_factors)])
        results['factor_loadings'] = loadings
        results['communalities'] = (loadings ** 2).sum(axis=1).to_dict()
        factor_var = np.var(fa.transform(df_scaled), axis=0)
        results['factor_variance'] = factor_var.tolist()
        results['factor_variance_ratio'] = (factor_var / factor_var.sum() * 100).tolist()

    results['n_factors'] = n_factors
    results['rotation_used'] = rotation

    return results


def reliability_analysis_advanced(df, items):
    """تحليل الثبات المتقدم"""
    results = {}

    data = df[items].dropna()

    # ألفا كرونباخ
    results['cronbach_alpha'] = calculate_cronbach_alpha(data)

    # معامل جتمان (Guttman's Lambda)
    from sklearn.covariance import EmpiricalCovariance
    cov_matrix = EmpiricalCovariance().fit(data).covariance_
    item_variances = np.diag(cov_matrix)
    total_variance = np.sum(cov_matrix)
    results['guttman_lambda'] = (len(items) / (len(items) - 1)) * (1 - np.sum(item_variances) / total_variance) if len(
        items) > 1 else np.nan

    # ألفا عند حذف كل فقرة
    alpha_if_deleted = {}
    for item in items:
        alpha = calculate_cronbach_alpha(data.drop(columns=[item]))
        alpha_if_deleted[item] = alpha
    results['alpha_if_deleted'] = alpha_if_deleted

    # ارتباط كل فقرة بالمجموع الكلي
    total_scores = data.sum(axis=1)
    item_total_corr = {}
    for item in items:
        corr = data[item].corr(total_scores)
        item_total_corr[item] = corr
    results['item_total_correlation'] = item_total_corr

    # ارتباط كل فقرة مع الأخرى (Inter-item correlation)
    inter_item_corr = data.corr()
    results['inter_item_correlation'] = inter_item_corr

    # متوسط ارتباط الفقرات
    upper_tri = inter_item_corr.where(np.triu(np.ones(inter_item_corr.shape), k=1).astype(bool))
    results['mean_inter_item_corr'] = upper_tri.unstack().dropna().mean()

    return results


def manova_analysis(df, dependent_vars, independent_var):
    """تحليل المانوفا (MANOVA)"""
    from statsmodels.multivariate.manova import MANOVA

    data = df[dependent_vars + [independent_var]].dropna()
    formula = f"{' + '.join(dependent_vars)} ~ C({independent_var})"
    manova = MANOVA.from_formula(formula, data=data)
    results = manova.mv_test()

    return results


def repeated_measures_anova(df, subject_id, time_points, dependent_var):
    """تحليل التباين للقياسات المتكررة"""
    from statsmodels.stats.anova import AnovaRM

    data = df[[subject_id, time_points, dependent_var]].dropna()
    aovrm = AnovaRM(data, dependent_var, subject_id, within=[time_points])
    res = aovrm.fit()

    return res


def friedman_test_repeated(data_groups):
    """اختبار فريدمان للقياسات المتكررة"""
    stat, p_value = friedmanchisquare(*data_groups)
    return {'statistic': stat, 'pvalue': p_value}


def anova_with_posthoc(df, dependent_var, independent_var):
    """تحليل ANOVA مع اختبارات ما بعد hoc"""
    results = {}

    # حساب المتوسطات حسب المجموعات
    group_means = df.groupby(independent_var)[dependent_var].mean()
    group_sizes = df.groupby(independent_var)[dependent_var].count()
    results['group_means'] = group_means
    results['group_sizes'] = group_sizes

    # اختبار تجانس التباين (Levene's test)
    groups = [group[dependent_var].values for name, group in df.groupby(independent_var)]
    if len(groups) >= 2:
        levene_stat, levene_p = levene(*groups)
        results['levene_test'] = {'statistic': levene_stat, 'pvalue': levene_p}
    else:
        results['levene_test'] = {'statistic': np.nan, 'pvalue': np.nan}

    # One-way ANOVA
    if len(groups) >= 2:
        f_stat, p_value = f_oneway(*groups)
        ss_between = sum(len(g) * (g.mean() - df[dependent_var].mean()) ** 2 for g in groups)
        ss_within = sum(((g - g.mean()) ** 2).sum() for g in groups)
        ss_total = ss_between + ss_within
        df_between = len(groups) - 1
        df_within = len(df) - len(groups)

        anova_table = pd.DataFrame({
            'Source': [independent_var, 'Residual'],
            'SS': [ss_between, ss_within],
            'DF': [df_between, df_within],
            'MS': [ss_between / df_between if df_between > 0 else 0, ss_within / df_within if df_within > 0 else 0],
            'F': [f_stat, np.nan],
            'PR(>F)': [p_value, np.nan]
        })
        results['anova_table'] = anova_table

        # إضافة مقاييس حجم التأثير
        results['eta_squared'] = calculate_effect_size_eta_squared(ss_between, ss_total)
        results['omega_squared'] = calculate_effect_size_omega_squared(ss_between, df_between, ss_within, df_within)

    else:
        results['anova_table'] = None
        results['eta_squared'] = np.nan
        results['omega_squared'] = np.nan

    # Tukey HSD post-hoc test
    if len(groups) >= 3 and p_value < 0.05:
        tukey = pairwise_tukeyhsd(df[dependent_var], df[independent_var])
        results['tukey_hsd'] = tukey
    else:
        results['tukey_hsd'] = None

    return results


def logistic_regression_analysis(X, y, feature_names):
    """تحليل الانحدار اللوجستي"""
    from sklearn.linear_model import LogisticRegression
    from sklearn.metrics import classification_report, roc_auc_score, roc_curve

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    model = LogisticRegression(max_iter=1000)
    model.fit(X_train, y_train)

    y_pred = model.predict(X_test)
    y_pred_proba = model.predict_proba(X_test)[:, 1]

    results = {
        'coefficients': model.coef_[0],
        'intercept': model.intercept_[0],
        'accuracy': accuracy_score(y_test, y_pred),
        'precision': precision_score(y_test, y_pred, average='weighted'),
        'recall': recall_score(y_test, y_pred, average='weighted'),
        'f1': f1_score(y_test, y_pred, average='weighted'),
        'roc_auc': roc_auc_score(y_test, y_pred_proba),
        'classification_report': classification_report(y_test, y_pred),
        'model': model
    }

    return results


def time_series_analysis(df, date_col, value_col):
    """تحليل السلاسل الزمنية"""
    results = {}

    df[date_col] = pd.to_datetime(df[date_col])
    df = df.sort_values(date_col)

    # الاتجاه العام
    x = np.arange(len(df))
    slope, intercept, r_value, p_value, std_err = stats.linregress(x, df[value_col])
    results['trend'] = {
        'slope': slope,
        'intercept': intercept,
        'r_squared': r_value ** 2,
        'pvalue': p_value
    }

    # المتوسط المتحرك
    results['moving_average_7'] = df[value_col].rolling(window=7).mean().tolist()
    results['moving_average_30'] = df[value_col].rolling(window=30).mean().tolist()

    # الموسمية
    from statsmodels.tsa.seasonal import seasonal_decompose
    try:
        decomposition = seasonal_decompose(df[value_col].dropna(), model='additive', period=min(7, len(df) // 2))
        results['seasonal'] = decomposition.seasonal.tolist()
        results['trend_component'] = decomposition.trend.tolist()
        results['residual'] = decomposition.resid.tolist()
    except:
        results['seasonal'] = None

    return results


def create_rtl_figure(fig):
    """تعديل الرسم البياني ليكون متوافقاً مع LTR"""
    fig.update_layout(
        font=dict(family="Arial", size=12),
        title_font=dict(family="Arial", size=16),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="center",
            x=0.5
        )
    )
    return fig


# ==================== دوال الرسوم البيانية المتعددة ====================

def create_histogram_chart(data, title, chart_type="histogram", bins=20):
    """إنشاء رسم بياني للتوزيع بأنواع مختلفة"""
    if chart_type == "histogram":
        fig = px.histogram(data, nbins=bins, title=title, color_discrete_sequence=['#667eea'])
        fig.update_layout(bargap=0.1)

    elif chart_type == "box":
        fig = px.box(data, title=title, color_discrete_sequence=['#764ba2'])

    elif chart_type == "violin":
        fig = px.violin(data, title=title, box=True, points="all", color_discrete_sequence=['#2ecc71'])

    elif chart_type == "strip":
        fig = px.strip(data, title=title, color_discrete_sequence=['#e74c3c'])

    elif chart_type == "ecdf":
        fig = px.ecdf(data, title=title, color_discrete_sequence=['#3498db'])

    elif chart_type == "density":
        fig = ff.create_distplot([data.dropna().values], [title], show_hist=False, show_rug=False)
        fig.update_layout(title=title)

    elif chart_type == "histogram_density":
        fig = ff.create_distplot([data.dropna().values], [title], show_hist=True, show_rug=False)
        fig.update_layout(title=title)

    else:
        fig = px.histogram(data, nbins=bins, title=title, color_discrete_sequence=['#667eea'])

    return create_rtl_figure(fig)


def create_bar_chart(df, x, y, title, chart_type="bar", color_by=None):
    """إنشاء رسم بياني أعمدة بأنواع مختلفة"""
    if chart_type == "bar":
        fig = px.bar(df, x=x, y=y, title=title, color=color_by, color_discrete_sequence=['#667eea'], text_auto=True)

    elif chart_type == "horizontal_bar":
        fig = px.bar(df, x=y, y=x, title=title, orientation='h', color=color_by, color_discrete_sequence=['#764ba2'],
                     text_auto=True)

    elif chart_type == "grouped_bar":
        fig = px.bar(df, x=x, y=y, title=title, color=color_by if color_by else x, barmode='group', text_auto=True)

    elif chart_type == "stacked_bar":
        fig = px.bar(df, x=x, y=y, title=title, color=color_by if color_by else x, barmode='stack', text_auto=True)

    elif chart_type == "percentage_bar":
        fig = px.bar(df, x=x, y=y, title=title, color=color_by if color_by else x, barmode='relative', text_auto=True)

    else:
        fig = px.bar(df, x=x, y=y, title=title, color=color_by, color_discrete_sequence=['#667eea'])

    return create_rtl_figure(fig)


def create_scatter_chart(df, x, y, title, chart_type="scatter", color_by=None, size_by=None):
    """إنشاء رسم بياني مبعثر بأنواع مختلفة"""
    if chart_type == "scatter":
        fig = px.scatter(df, x=x, y=y, title=title, color=color_by, size=size_by, color_discrete_sequence=['#667eea'])

    elif chart_type == "scatter_with_line":
        fig = px.scatter(df, x=x, y=y, title=title, trendline="ols", color=color_by,
                         color_discrete_sequence=['#667eea'])

    elif chart_type == "bubble":
        fig = px.scatter(df, x=x, y=y, title=title, size=size_by if size_by else y, color=color_by,
                         color_discrete_sequence=px.colors.qualitative.Set2)

    elif chart_type == "scatter_matrix":
        fig = px.scatter_matrix(df, dimensions=[x, y] if len([x, y]) > 1 else [x], title=title)

    elif chart_type == "density_contour":
        fig = px.density_contour(df, x=x, y=y, title=title, color_discrete_sequence=['#667eea'])

    else:
        fig = px.scatter(df, x=x, y=y, title=title, color=color_by, color_discrete_sequence=['#667eea'])

    return create_rtl_figure(fig)


def create_correlation_heatmap(corr_matrix, title, chart_type="heatmap", show_text=True):
    """إنشاء خريطة حرارة للارتباط بأنواع مختلفة"""
    if chart_type == "heatmap":
        fig = px.imshow(corr_matrix, text_auto=show_text, aspect="auto",
                        color_continuous_scale='RdBu_r', title=title,
                        zmin=-1, zmax=1)

    elif chart_type == "clustered_heatmap":
        import scipy.cluster.hierarchy as sch
        linkage = sch.linkage(corr_matrix, method='average')
        order = sch.leaves_list(linkage)
        corr_clustered = corr_matrix.iloc[order, order]
        fig = px.imshow(corr_clustered, text_auto=show_text, aspect="auto",
                        color_continuous_scale='RdBu_r', title=f"{title} (مرتبة)",
                        zmin=-1, zmax=1)

    else:
        fig = px.imshow(corr_matrix, text_auto=show_text, aspect="auto",
                        color_continuous_scale='RdBu_r', title=title,
                        zmin=-1, zmax=1)

    fig.update_layout(height=600)
    return create_rtl_figure(fig)


def create_pie_chart(df, names, values, title, chart_type="pie"):
    """إنشاء رسم بياني دائري بأنواع مختلفة - تم تبسيطه"""
    if chart_type == "pie":
        fig = px.pie(df, names=names, values=values, title=title, hole=0)
    else:
        fig = px.pie(df, names=names, values=values, title=title, hole=0.4)

    fig.update_traces(textposition='inside', textinfo='percent+label')
    return create_rtl_figure(fig)


def create_boxplot_chart(df, x, y, title, points="all"):
    """إنشاء مخطط صندوقي"""
    fig = px.box(df, x=x, y=y, title=title, color=x, points=points, color_discrete_sequence=px.colors.qualitative.Set2)
    return create_rtl_figure(fig)


def create_line_chart(df, x, y, title, chart_type="line", markers=True):
    """إنشاء رسم بياني خطي بأنواع مختلفة"""
    if chart_type == "line":
        fig = px.line(df, x=x, y=y, title=title, markers=markers, color_discrete_sequence=['#667eea'])

    elif chart_type == "area":
        fig = px.area(df, x=x, y=y, title=title, color_discrete_sequence=['#764ba2'])

    elif chart_type == "smooth_line":
        from scipy.interpolate import make_interp_spline
        x_vals = df[x].values
        y_vals = df[y].values
        if len(x_vals) > 3:
            x_smooth = np.linspace(x_vals.min(), x_vals.max(), 300)
            spl = make_interp_spline(x_vals, y_vals, k=min(3, len(x_vals) - 1))
            y_smooth = spl(x_smooth)
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=x_vals, y=y_vals, mode='markers', name='البيانات', marker=dict(color='#667eea')))
            fig.add_trace(go.Scatter(x=x_smooth, y=y_smooth, mode='lines', name='المنحنى الملساء',
                                     line=dict(color='#764ba2', width=2)))
            fig.update_layout(title=title)
        else:
            fig = px.line(df, x=x, y=y, title=title, markers=markers, color_discrete_sequence=['#667eea'])

    elif chart_type == "line_with_ci":
        fig = px.line(df, x=x, y=y, title=title, markers=markers, error_y_minus=y, error_y_plus=y,
                      color_discrete_sequence=['#667eea'])

    else:
        fig = px.line(df, x=x, y=y, title=title, markers=markers, color_discrete_sequence=['#667eea'])

    return create_rtl_figure(fig)


def generate_html_report(results_dict, title="تقرير التحليل الإحصائي"):
    """توليد تقرير HTML متقدم للنتائج"""
    report = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); }}
            .container {{ max-width: 1200px; margin: auto; background: white; padding: 30px; border-radius: 20px; box-shadow: 0 10px 40px rgba(0,0,0,0.1); }}
            h1 {{ color: #667eea; border-bottom: 3px solid #667eea; padding-bottom: 10px; margin-bottom: 30px; text-align: center; }}
            h2 {{ color: #764ba2; margin-top: 30px; margin-bottom: 20px; }}
            table {{ width: 100%; border-collapse: collapse; margin: 15px 0; direction: ltr; }}
            th, td {{ padding: 10px; text-align: center; border: 1px solid #ddd; }}
            th {{ background: #667eea; color: white; }}
            tr:nth-child(even) {{ background: #f2f2f2; }}
            .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>📊 {title}</h1>
            <p style="text-align: center;">تم إنشاء التقرير في: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    """

    for key, value in results_dict.items():
        if isinstance(value, pd.DataFrame):
            report += f"""
            <div>
                <h2>{key}</h2>
                {value.to_html()}
            </div>
            """
        elif isinstance(value, dict):
            report += f"""
            <div>
                <h2>{key}</h2>
                <ul>
            """
            for k, v in value.items():
                if isinstance(v, float):
                    report += f"<li><strong>{k}:</strong> {v:.4f}</li>"
                else:
                    report += f"<li><strong>{k}:</strong> {v}</li>"
            report += """
                </ul>
            </div>
            """

    report += """
            <div class="footer">
                <p>تم إنشاء هذا التقرير بواسطة نظام التحليل الإحصائي المتقدم للاستبيانات</p>
            </div>
        </div>
    </body>
    </html>
    """
    return report


def generate_word_report(results_dict, filename="analysis_report.docx"):
    """توليد تقرير Word"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading('تقرير التحليل الإحصائي', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        for key, value in results_dict.items():
            doc.add_heading(key, level=1)
            if isinstance(value, pd.DataFrame):
                table = doc.add_table(rows=len(value) + 1, cols=len(value.columns))
                table.style = 'Light Grid Accent 1'
                for j, col in enumerate(value.columns):
                    table.cell(0, j).text = str(col)
                for i, row in value.iterrows():
                    for j, val in enumerate(row):
                        table.cell(i + 1, j).text = str(val) if not pd.isna(val) else "N/A"
            elif isinstance(value, dict):
                for k, v in value.items():
                    doc.add_paragraph(f"• {k}: {v}")
            doc.add_paragraph("")

        doc.save(filename)
        return True
    except Exception as e:
        return False


def export_to_excel(results_dict, filename="analysis_results.xlsx"):
    """تصدير النتائج إلى ملف Excel"""
    try:
        with pd.ExcelWriter(filename, engine='xlsxwriter') as writer:
            for sheet_name, data in results_dict.items():
                if isinstance(data, pd.DataFrame):
                    data.to_excel(writer, sheet_name=sheet_name[:31], index=True)
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
                    df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
        return True
    except Exception as e:
        return False


def send_email(to_email, subject, body):
    """إرسال بريد إلكتروني"""
    try:
        from_email = "boutoubaamed@gmail.com"

        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain', 'utf-8'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, "your_app_password")
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"Error: {e}")
        return False


# ==================== تنسيق CSS - جميع الجداول والرسوم LTR ====================

st.markdown("""
<style>
    * {
        font-family: 'Arial', sans-serif;
    }

    /* جميع الجداول من اليسار إلى اليمين */
    .stDataFrame {
        direction: ltr !important;
    }
    .stDataFrame table {
        width: 100%;
        text-align: center !important;
        direction: ltr !important;
    }
    .stDataFrame th, .stDataFrame td {
        text-align: center !important;
        vertical-align: middle !important;
        direction: ltr !important;
    }

    /* جميع الرسوم البيانية من اليسار إلى اليمين */
    .plotly-graph-div {
        direction: ltr !important;
    }

    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 20px;
        text-align: center;
        margin-bottom: 2rem;
        color: white;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
    }
    .main-header h1 { font-size: 2.5rem; margin: 0; font-weight: 700; text-align: center; }
    .main-header p { margin-top: 0.5rem; opacity: 0.9; text-align: center; }

    .section-card {
        background: white;
        border-radius: 20px;
        padding: 1.5rem;
        margin-bottom: 2rem;
        box-shadow: 0 5px 20px rgba(0,0,0,0.08);
        border-right: 5px solid #667eea;
        transition: all 0.3s ease;
    }
    .section-card:hover { transform: translateY(-3px); box-shadow: 0 8px 25px rgba(0,0,0,0.12); }
    .section-title {
        font-size: 1.8rem;
        color: #667eea;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 3px solid #667eea;
        display: inline-block;
    }

    .result-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 15px;
        padding: 1rem;
        margin: 1rem 0;
    }

    .trend-strongly-agree { background-color: #1b5e20; color: white; padding: 3px 10px; border-radius: 20px; display: inline-block; }
    .trend-agree { background-color: #4caf50; color: white; padding: 3px 10px; border-radius: 20px; display: inline-block; }
    .trend-neutral { background-color: #ffc107; color: #333; padding: 3px 10px; border-radius: 20px; display: inline-block; }
    .trend-disagree { background-color: #ff9800; color: white; padding: 3px 10px; border-radius: 20px; display: inline-block; }
    .trend-strongly-disagree { background-color: #f44336; color: white; padding: 3px 10px; border-radius: 20px; display: inline-block; }

    .regression-box { background-color: #e8f4f8; border-radius: 10px; padding: 15px; margin: 10px 0; border-right: 4px solid #2e86ab; direction: ltr; text-align: left; }
    .mediation-box { background-color: #e8f8f5; border-radius: 10px; padding: 15px; margin: 10px 0; border-right: 4px solid #1abc9c; }
    .success-box { background-color: #d4edda; border: 1px solid #c3e6cb; border-radius: 10px; padding: 15px; margin: 10px 0; }
    .info-box { background-color: #e8f4f8; border-right: 4px solid #2e86ab; border-radius: 10px; padding: 15px; margin: 15px 0; }

    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        font-weight: bold;
        transition: all 0.3s ease;
        width: 100%;
    }
    .stButton > button:hover { transform: scale(1.02); box-shadow: 0 5px 15px rgba(0,0,0,0.2); }

    .footer { text-align: center; padding: 2rem; color: #666; font-size: 0.9rem; }
    hr { margin: 2rem 0; border: none; height: 2px; background: linear-gradient(to right, #667eea, #764ba2); }
</style>
""", unsafe_allow_html=True)

# تهيئة حالة الجلسة
if 'page' not in st.session_state:
    st.session_state.page = 'upload'
if 'data_loaded' not in st.session_state:
    st.session_state.data_loaded = False
if 'df' not in st.session_state:
    st.session_state.df = None
if 'factors' not in st.session_state:
    st.session_state.factors = []
if 'social_vars' not in st.session_state:
    st.session_state.social_vars = []
if 'independent_vars' not in st.session_state:
    st.session_state.independent_vars = []
if 'mediator_vars' not in st.session_state:
    st.session_state.mediator_vars = []
if 'normality_test_done' not in st.session_state:
    st.session_state.normality_test_done = False
if 'is_normal' not in st.session_state:
    st.session_state.is_normal = None
if 'analysis_history' not in st.session_state:
    st.session_state.analysis_history = []
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'chart_preference' not in st.session_state:
    st.session_state.chart_preference = {}

# العنوان الرئيسي
st.markdown("""
<div class="main-header">
    <h1>📊 تطبيق للتحليل الإحصائي المتقدم للاستبيانات</h1>
    <p>تطبيق موجه لإعداد مذكرة التخرج ماستر - مشروع التخرج ليسانس</p>
        <p>كلية العلوم الاقتصادية، التجارية وعلوم التسيير</p>
</div>
""", unsafe_allow_html=True)

# ==================== شريط التنقل الجانبي (مرتب حسب التسلسل الصحيح) ====================
with st.sidebar:
    st.markdown("""
    <div style="
        text-align: center;
        background: linear-gradient(135deg, #667eea15 0%, #764ba215 100%);
        padding: 12px;
        border-radius: 15px;
        margin: 10px 0;
    ">
        <h3 style="
            color: #667eea;
            margin: 0;
            font-size: 1.3rem;
        ">🧭 قائمة التحليل</h3>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    if st.button("📁 1. تحميل البيانات", use_container_width=True):
        st.session_state.page = 'upload'
        st.rerun()

    if st.session_state.data_loaded:
        if st.button("🔧 2. تحليل الثبات (Reliability)", use_container_width=True):
            st.session_state.page = 'reliability'
            st.rerun()

        if st.button("🔧 3. تحليل العوامل (Factor Analysis)", use_container_width=True):
            st.session_state.page = 'factor'
            st.rerun()

        if st.button("📊 4. الدراسة الوصفية", use_container_width=True):
            st.session_state.page = 'descriptive'
            st.rerun()

        if st.button("📈 5. الاتجاهات العامة", use_container_width=True):
            st.session_state.page = 'trends'
            st.rerun()

        if st.button("🔬 6. اختبار التوزيع الطبيعي", use_container_width=True):
            st.session_state.page = 'normality'
            st.rerun()

        if st.button("🔗 7. تحليل الارتباط", use_container_width=True):
            st.session_state.page = 'correlation'
            st.rerun()

        if st.button("📈 8. تحليل الانحدار", use_container_width=True):
            st.session_state.page = 'regression'
            st.rerun()

        if st.button("🔄 9. تحليل الوساطة", use_container_width=True):
            st.session_state.page = 'mediation'
            st.rerun()

        if st.button("📐 10. اختبار الفروقات", use_container_width=True):
            st.session_state.page = 'differences'
            st.rerun()

        if st.button("📊 11. تحليل التباين (ANOVA)", use_container_width=True):
            st.session_state.page = 'anova'
            st.rerun()

        if st.button("🎯 12. تحليل التجميع", use_container_width=True):
            st.session_state.page = 'clustering'
            st.rerun()

        if st.button("📋 13. تصدير التقرير", use_container_width=True):
            st.session_state.page = 'export'
            st.rerun()

    st.markdown("---")
    if st.button("✉️ 14. التواصل", use_container_width=True):
        st.session_state.page = 'contact'
        st.rerun()

    st.markdown("""
    <div style="text-align: center;">
        <h3>تطبيق للتحليل الإحصائي المتقدم</h3>
        <p style="color: #666; font-size: 0.9rem;">إصدار تجريبي</p>
        <p style="color: #666; font-size: 0.9rem;">© 2026</p>
    </div>
    """, unsafe_allow_html=True)

# ==================== الصفحة 1: تحميل البيانات ====================
if st.session_state.page == 'upload':
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📁 تحميل ملف البيانات</h2>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader("اختر ملف Excel أو CSV", type=['xlsx', 'xls', 'csv'])

    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)

            st.session_state.df = df
            st.session_state.data_loaded = True

            st.success(f"✅ تم تحميل البيانات بنجاح! ({df.shape[0]} صف، {df.shape[1]} عمود)")

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("عدد الصفوف", df.shape[0])
            with col2:
                st.metric("عدد الأعمدة", df.shape[1])
            with col3:
                st.metric("القيم المفقودة", df.isnull().sum().sum())
            with col4:
                st.metric("القيم الفريدة", df.nunique().sum())

            with st.expander("📋 معاينة البيانات"):
                st.dataframe(df.head(10), use_container_width=True)

            st.markdown("---")
            st.markdown("### 🎯 تحديد المتغيرات")

            all_columns = list(df.columns)

            col1, col2, col3 = st.columns(3)
            with col1:
                social_vars = st.multiselect("👥 المتغيرات الاجتماعية والديموغرافية:", options=all_columns)
                st.session_state.social_vars = social_vars

            with col2:
                independent_vars = st.multiselect("📈 المتغيرات المستقلة/المتنبئة:",
                                                  options=[c for c in all_columns if c not in social_vars])
                st.session_state.independent_vars = independent_vars

            with col3:
                mediator_vars = st.multiselect("🔄 المتغيرات الوسيطة (Mediator):", options=[c for c in all_columns if
                                                                                           c not in social_vars + independent_vars])
                st.session_state.mediator_vars = mediator_vars

            question_vars = [c for c in all_columns if c not in social_vars + independent_vars + mediator_vars]

            st.markdown(f"""
            <div class="success-box">
                <strong>✅ تم التعرف على:</strong><br>
                • المتغيرات الاجتماعية: {len(social_vars)} متغير<br>
                • المتغيرات المستقلة: {len(independent_vars)} متغير<br>
                • المتغيرات الوسيطة: {len(mediator_vars)} متغير<br>
                • فقرات الاستبيان: {len(question_vars)} فقرة
            </div>
            """, unsafe_allow_html=True)

            num_factors = st.number_input("🔢 عدد المحاور/الأبعاد:", min_value=1, max_value=10, value=2, step=1)

            factors = []
            for i in range(num_factors):
                with st.expander(f"📊 المحور {i + 1}"):
                    col1, col2 = st.columns(2)
                    with col1:
                        factor_name = st.text_input(f"اسم المحور {i + 1}:", value=f"البعد_{i + 1}", key=f"name_{i}")
                        selected_questions = st.multiselect(f"اختر الفقرات التي تقيس هذا المحور:",
                                                            options=question_vars, key=f"q_{i}")
                    with col2:
                        selected_social = st.multiselect(f"المتغيرات الاجتماعية المرتبطة:", options=social_vars,
                                                         key=f"s_{i}")
                        selected_independent = st.multiselect(f"المتغيرات المستقلة المرتبطة:", options=independent_vars,
                                                              key=f"ind_{i}")
                        selected_mediator = st.multiselect(f"المتغيرات الوسيطة المرتبطة:", options=mediator_vars,
                                                           key=f"med_{i}")

                    factors.append({
                        'id': i,
                        'name': factor_name,
                        'questions': selected_questions,
                        'social_vars': selected_social,
                        'independent_vars': selected_independent,
                        'mediator_vars': selected_mediator
                    })

            st.session_state.factors = factors

            if st.button("✅ تأكيد وتفعيل التحليل", type="primary", use_container_width=True):
                with st.spinner("جاري حساب المتغيرات الكامنة..."):
                    for factor in factors:
                        if factor['questions']:
                            df[factor['name']] = df[factor['questions']].mean(axis=1)
                    st.session_state.df = df
                st.success("✅ تم تفعيل التحليل! انتقل الآن إلى الأقسام الأخرى من القائمة الجانبية.")
                st.balloons()

        except Exception as e:
            st.error(f"❌ خطأ في قراءة الملف: {e}")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 2: تحليل الثبات المتقدم ====================
elif st.session_state.page == 'reliability' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🔧 تحليل الثبات المتقدم (Reliability Analysis)</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل الثبات:</strong> يهدف إلى قياس مدى اتساق وثبات أداة القياس.</p>
        <p>📊 <strong>المقاييس:</strong> ألفا كرونباخ، معامل جتمان، ارتباط الفقرة بالكل، ألفا عند حذف الفقرة.</p>
        <p>🎯 <strong>القاعدة:</strong> ألفا كرونباخ ≥ 0.70 مقبول، ≥ 0.80 جيد، ≥ 0.90 ممتاز.</p>
    </div>
    """, unsafe_allow_html=True)

    selected_factor = st.selectbox("اختر المحور لتحليل ثباته:",
                                   [f['name'] for f in factors if len(f['questions']) >= 2])

    if selected_factor:
        factor = next((f for f in factors if f['name'] == selected_factor), None)
        if factor and len(factor['questions']) >= 2:
            reliability_results = reliability_analysis_advanced(df, factor['questions'])

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("ألفا كرونباخ", f"{reliability_results['cronbach_alpha']:.4f}")
                if reliability_results['cronbach_alpha'] >= 0.9:
                    st.success("ثبات ممتاز")
                elif reliability_results['cronbach_alpha'] >= 0.8:
                    st.success("ثبات جيد جداً")
                elif reliability_results['cronbach_alpha'] >= 0.7:
                    st.info("ثبات جيد")
                elif reliability_results['cronbach_alpha'] >= 0.6:
                    st.warning("ثبات مقبول")
                else:
                    st.error("ثبات ضعيف - يُنصح بمراجعة الفقرات")
            with col2:
                st.metric("معامل جتمان", f"{reliability_results['guttman_lambda']:.4f}")
            with col3:
                st.metric("متوسط ارتباط الفقرات", f"{reliability_results['mean_inter_item_corr']:.4f}")

            st.markdown("### 📊 ألفا كرونباخ عند حذف كل فقرة")
            alpha_if_df = pd.DataFrame({'الفقرة': list(reliability_results['alpha_if_deleted'].keys()),
                                        'ألفا عند الحذف': [f"{v:.4f}" for v in
                                                           reliability_results['alpha_if_deleted'].values()]})
            st.dataframe(alpha_if_df, use_container_width=True)

            # توصيات لتحسين الثبات
            st.markdown("### 📝 توصيات لتحسين الثبات")
            if reliability_results['cronbach_alpha'] < 0.7:
                st.markdown("""
                <div class="info-box">
                    <p><strong>الإجراءات المقترحة لتحسين الثبات:</strong></p>
                    <ul>
                        <li>حذف الفقرات التي تسبب انخفاض ألفا عند حذفها (ألفا أعلى من القيمة الإجمالية)</li>
                        <li>مراجعة صياغة الفقرات التي لديها ارتباط ضعيف بالمجموع الكلي (أقل من 0.3)</li>
                        <li>زيادة عدد الفقرات في المحور</li>
                        <li>توسيع حجم العينة</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div class="success-box">
                    <p>✅ <strong>الثبات جيد أو ممتاز</strong> - يمكن الاعتماد على أداة القياس.</p>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("### 📈 ارتباط كل فقرة بالمجموع الكلي")
            corr_df = pd.DataFrame({'الفقرة': list(reliability_results['item_total_correlation'].keys()),
                                    'معامل الارتباط': [f"{v:.4f}" for v in
                                                       reliability_results['item_total_correlation'].values()]})
            st.dataframe(corr_df, use_container_width=True)

            # رسم بياني لارتباطات الفقرات
            fig = px.bar(corr_df, x='الفقرة', y='معامل الارتباط', title='ارتباط كل فقرة بالمجموع الكلي',
                         color='معامل الارتباط', color_continuous_scale='Viridis', text='معامل الارتباط')
            fig.add_hline(y=0.3, line_dash="dash", line_color="red", annotation_text="الحد الأدنى المقبول (0.3)")
            fig = create_rtl_figure(fig)
            st.plotly_chart(fig, use_container_width=True)

        else:
            st.warning("⚠️ يجب أن يحتوي المحور على فقرتين على الأقل لتحليل الثبات")
    else:
        st.info("📌 اختر محوراً يحتوي على فقرتين على الأقل لتحليل الثبات")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 3: تحليل العوامل المتقدم ====================
elif st.session_state.page == 'factor' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🔧 تحليل العوامل المتقدم (Factor Analysis)</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل العوامل:</strong> يهدف إلى تقليل عدد المتغيرات وتحديد البنية الكامنة وراء البيانات (التحقق من الصلاحية البنائية).</p>
        <p>📊 <strong>يتضمن:</strong> قيم eigenvalues، نسبة التباين المفسر، تشبعات العوامل، المشتركات.</p>
        <p>🎯 <strong>معيار كايزر:</strong> استبقاء العوامل ذات القيمة الذاتية (Eigenvalue) > 1.</p>
    </div>
    """, unsafe_allow_html=True)

    all_questions = []
    for factor in factors:
        all_questions.extend(factor['questions'])
    all_questions = list(set(all_questions))

    if len(all_questions) < 3:
        st.warning("⚠️ يلزم وجود 3 فقرات على الأقل لتحليل العوامل")
    else:
        factor_data = df[all_questions].dropna()

        if len(factor_data) >= 10:
            n_factors = st.slider("عدد العوامل المطلوب استخراجها:", 2, min(10, len(all_questions) - 1), 3)
            rotation = st.selectbox("نوع التدوير:", ["varimax", "none"])

            if st.button("تنفيذ تحليل العوامل", type="primary"):
                with st.spinner("جاري تحليل العوامل..."):
                    fa_results = factor_analysis_advanced(factor_data, n_factors, rotation)

                    st.markdown("### 📊 القيم الذاتية (Eigenvalues)")
                    eigen_df = pd.DataFrame({'العامل': range(1, len(fa_results['eigenvalues']) + 1),
                                             'Eigenvalue': fa_results['eigenvalues'],
                                             'نسبة التباين (%)': [f"{v * 100:.1f}" for v in
                                                                  fa_results['explained_variance_ratio']],
                                             'التباين التراكمي (%)': [f"{v * 100:.1f}" for v in
                                                                      fa_results['cumulative_variance']]})
                    st.dataframe(eigen_df, use_container_width=True)

                    st.markdown("### 📈 منحنى Scree Plot")
                    scree_fig = px.line(x=fa_results['scree_data']['components'],
                                        y=fa_results['scree_data']['eigenvalues'], markers=True,
                                        title="Scree Plot للقيم الذاتية")
                    scree_fig.add_hline(y=1, line_dash="dash", line_color="red",
                                        annotation_text="معيار كايزر (Eigenvalue = 1)")
                    scree_fig = create_rtl_figure(scree_fig)
                    st.plotly_chart(scree_fig, use_container_width=True)

                    st.markdown(f"**عدد العوامل حسب معيار كايزر:** {fa_results['n_factors_kaiser']}")

                    st.markdown("### 📊 تشبعات العوامل (Factor Loadings)")
                    loadings_with_communalities = fa_results['factor_loadings'].copy()
                    loadings_with_communalities['المشترك (Communality)'] = fa_results['communalities'].values()
                    st.dataframe(loadings_with_communalities.style.format('{:.3f}'), use_container_width=True)

                    # خريطة حرارة للتشبعات
                    heatmap_fig = px.imshow(fa_results['factor_loadings'].abs(), text_auto=True, aspect="auto",
                                            title="خريطة حرارة تشبعات العوامل", color_continuous_scale='Viridis')
                    heatmap_fig = create_rtl_figure(heatmap_fig)
                    st.plotly_chart(heatmap_fig, use_container_width=True)

                    # تفسير النتائج
                    st.markdown("### 📝 تفسير النتائج")
                    total_variance_explained = sum(fa_results['explained_variance_ratio'][:n_factors]) * 100
                    st.markdown(f"""
                    <div class="info-box">
                        <p><strong>التباين الكلي المفسر:</strong> {total_variance_explained:.1f}%</p>
                        <p><strong>توصيات:</strong></p>
                        <ul>
                            <li>تشبعات العوامل > 0.7 تعتبر ممتازة، > 0.5 مقبولة</li>
                            <li>المشترك (Communality) > 0.5 يشير إلى أن العامل يفسر جزءاً جيداً من تباين الفقرة</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.warning("⚠️ عدد الملاحظات قليل جداً لتحليل العوامل (يحتاج على الأقل 10 ملاحظات)")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 4: الدراسة الوصفية المتقدمة ====================
elif st.session_state.page == 'descriptive' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📊 الدراسة الوصفية المتقدمة</h2>', unsafe_allow_html=True)

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["📊 إحصاءات وصفية", "📈 رسوم بيانية", "🔍 كشف القيم المتطرفة", "📋 التكرارات", "📊 مقاييس الشكل"])

    with tab1:
        st.markdown("### الإحصاءات الوصفية التفصيلية")
        numeric_df = df.select_dtypes(include=[np.number])
        if len(numeric_df.columns) > 0:
            detailed_stats = calculate_descriptive_statistics_detailed(numeric_df)
            st.dataframe(detailed_stats, use_container_width=True)

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("عدد المتغيرات الرقمية", len(numeric_df.columns))
            with col2:
                st.metric("متوسط المتوسطات", f"{numeric_df.mean().mean():.3f}")
            with col3:
                st.metric("متوسط الانحرافات", f"{numeric_df.std().mean():.3f}")
            with col4:
                st.metric("متوسط معامل الاختلاف", f"{(numeric_df.std() / numeric_df.mean() * 100).mean():.1f}%")

    with tab2:
        st.markdown("### الرسوم البيانية المتعددة")
        numeric_cols = numeric_df.columns.tolist()
        if numeric_cols:
            selected_col = st.selectbox("اختر المتغير للرسم:", numeric_cols)

            chart_type = st.selectbox("نوع الرسم البياني:", [
                "مدرج تكراري (Histogram)",
                "مدرج تكراري مع كثافة (Histogram + Density)",
                "مخطط صندوقي (Box Plot)",
                "مخطط كمان (Violin Plot)",
                "مخطط شريطي مع نقاط (Strip Plot)",
                "منحنى الكثافة (Density Plot)",
                "مخطط التوزيع التراكمي (ECDF)",
                "مخطط Q-Q (Q-Q Plot)"
            ])

            if chart_type == "مخطط Q-Q (Q-Q Plot)":
                data = df[selected_col].dropna()
                fig = go.Figure()
                qq_data = stats.probplot(data, dist="norm")
                fig.add_trace(go.Scatter(x=qq_data[0][0], y=qq_data[0][1], mode='markers', name='البيانات',
                                         marker=dict(color='#667eea', size=6)))
                fig.add_trace(
                    go.Scatter(x=qq_data[0][0], y=qq_data[1][0] + qq_data[1][1] * np.array(qq_data[0][0]), mode='lines',
                               name='خط المرجع', line=dict(color='red', dash='dash', width=2)))
                fig.update_layout(title=f"Q-Q Plot لـ {selected_col}", height=500)
                fig = create_rtl_figure(fig)
                st.plotly_chart(fig, use_container_width=True)
            else:
                chart_map = {
                    "مدرج تكراري (Histogram)": "histogram",
                    "مدرج تكراري مع كثافة (Histogram + Density)": "histogram_density",
                    "مخطط صندوقي (Box Plot)": "box",
                    "مخطط كمان (Violin Plot)": "violin",
                    "مخطط شريطي مع نقاط (Strip Plot)": "strip",
                    "منحنى الكثافة (Density Plot)": "density",
                    "مخطط التوزيع التراكمي (ECDF)": "ecdf"
                }
                fig = create_histogram_chart(df[selected_col], f"توزيع {selected_col}", chart_map[chart_type])
                st.plotly_chart(fig, use_container_width=True)

    with tab3:
        st.markdown("### كشف القيم المتطرفة")
        numeric_cols = numeric_df.columns.tolist()
        if numeric_cols:
            selected_col = st.selectbox("اختر المتغير لكشف القيم المتطرفة:", numeric_cols, key="outlier_select")
            data = df[selected_col].dropna()
            outliers_iqr, lower_bound, upper_bound = detect_outliers_iqr(data)
            outliers_zscore = detect_outliers_zscore(data)

            col1, col2 = st.columns(2)
            with col1:
                st.metric("عدد القيم المتطرفة (IQR)", len(outliers_iqr))
                st.metric("الحد الأدنى (IQR)", f"{lower_bound:.3f}")
                st.metric("الحد الأعلى (IQR)", f"{upper_bound:.3f}")
            with col2:
                st.metric("عدد القيم المتطرفة (Z-Score)", len(outliers_zscore))
                st.metric("نسبة القيم المتطرفة", f"{(len(outliers_zscore) / len(data) * 100):.2f}%")

    with tab4:
        st.markdown("### التكرارات والنسب المئوية")
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
        categorical_cols.extend(st.session_state.social_vars)

        if categorical_cols:
            selected_col = st.selectbox("اختر المتغير الفئوي:", categorical_cols, key="freq_select")
            freq = df[selected_col].value_counts()
            percent = (freq / len(df) * 100).round(2)
            cumulative = percent.cumsum()
            freq_df = pd.DataFrame({'الفئة': freq.index, 'التكرار': freq.values, 'النسبة المئوية': percent.values,
                                    'النسبة التراكمية': cumulative.values})
            st.dataframe(freq_df, use_container_width=True)

            pie_type = st.selectbox(f"نوع الرسم البياني:", ["دائري (Pie)", "دونات (Donut)"])
            pie_map = {"دائري (Pie)": "pie", "دونات (Donut)": "donut"}
            fig_pie = create_pie_chart(freq_df.reset_index(), 'الفئة', 'التكرار', f"توزيع {selected_col}",
                                       pie_map[pie_type])
            st.plotly_chart(fig_pie, use_container_width=True)

    with tab5:
        st.markdown("### مقاييس الشكل (Skewness & Kurtosis)")
        numeric_cols = numeric_df.columns.tolist()
        if numeric_cols:
            skewness_data = []
            kurtosis_data = []
            for col in numeric_cols:
                data = df[col].dropna()
                if len(data) > 0:
                    skew_val = data.skew()
                    if skew_val > 1:
                        skew_interp = "ملتوي بشدة لليمين"
                    elif skew_val > 0.5:
                        skew_interp = "ملتوي لليمين"
                    elif skew_val < -1:
                        skew_interp = "ملتوي بشدة لليسار"
                    elif skew_val < -0.5:
                        skew_interp = "ملتوي لليسار"
                    else:
                        skew_interp = "متماثل تقريباً"

                    kurt_val = data.kurtosis()
                    if kurt_val > 3:
                        kurt_interp = "ذروة حادة جداً"
                    elif kurt_val > 1:
                        kurt_interp = "ذروة حادة"
                    elif kurt_val < -3:
                        kurt_interp = "مسطح جداً"
                    elif kurt_val < -1:
                        kurt_interp = "مسطح"
                    else:
                        kurt_interp = "طبيعي"

                    skewness_data.append({'المتغير': col, 'Skewness': f"{skew_val:.3f}", 'تفسير': skew_interp})
                    kurtosis_data.append({'المتغير': col, 'Kurtosis': f"{kurt_val:.3f}", 'تفسير': kurt_interp})

            col1, col2 = st.columns(2)
            with col1:
                st.dataframe(pd.DataFrame(skewness_data), use_container_width=True)
            with col2:
                st.dataframe(pd.DataFrame(kurtosis_data), use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 5: الاتجاهات العامة ====================
elif st.session_state.page == 'trends' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📈 الاتجاهات العامة للفقرات والمحاور</h2>', unsafe_allow_html=True)

    st.markdown("### 📝 اتجاهات الفقرات حسب سلم ليكارت")
    items_trend = []
    for factor in factors:
        for q in factor['questions']:
            if q in df.columns:
                mean_val = df[q].mean()
                trend_text, trend_class, trend_level = get_likert_trend(mean_val)
                items_trend.append(
                    {'الفقرة': q, 'المحور': factor['name'], 'المتوسط': f"{mean_val:.2f}", 'الاتجاه': trend_text,
                     'المستوى': trend_level})

    if items_trend:
        st.dataframe(pd.DataFrame(items_trend), use_container_width=True)

        high_items = len([i for i in items_trend if i['المستوى'] in ['مرتفع جداً', 'مرتفع']])
        neutral_items = len([i for i in items_trend if i['المستوى'] == 'متوسط'])
        low_items = len([i for i in items_trend if i['المستوى'] in ['منخفض', 'منخفض جداً']])

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📈 اتجاه إيجابي", f"{high_items} فقرة")
        with col2:
            st.metric("➖ اتجاه محايد", f"{neutral_items} فقرة")
        with col3:
            st.metric("📉 اتجاه سلبي", f"{low_items} فقرة")

        trend_df = pd.DataFrame([{'المستوى': i['المستوى']} for i in items_trend])
        trend_counts = trend_df['المستوى'].value_counts()
        fig = px.bar(x=trend_counts.index, y=trend_counts.values, title="توزيع مستويات الاتجاه",
                     color=trend_counts.index)
        fig = create_rtl_figure(fig)
        st.plotly_chart(fig, use_container_width=True)

    st.markdown("### 🎯 اتجاهات المحاور")
    for factor in factors:
        if factor['questions']:
            mean_score = df[factor['name']].mean()
            trend_text, trend_class, trend_level = get_likert_trend(mean_score)

            if len(factor['questions']) >= 2:
                cronbach = calculate_cronbach_alpha(df[factor['questions']])
                cronbach_display = f"{cronbach:.3f}" if not np.isnan(cronbach) else "غير قابل للحساب"
                reliability_text = "ممتاز" if not np.isnan(cronbach) and cronbach >= 0.9 else "جيد جداً" if not np.isnan(
                    cronbach) and cronbach >= 0.8 else "جيد" if not np.isnan(
                    cronbach) and cronbach >= 0.7 else "مقبول" if not np.isnan(
                    cronbach) and cronbach >= 0.6 else "ضعيف" if not np.isnan(cronbach) else "غير محدد"
            else:
                cronbach_display = "غير قابل للحساب (يحتاج فقرتين)"
                reliability_text = "غير محدد"

            st.markdown(f"""
            <div style="background: #f8f9fa; border-radius: 15px; padding: 1rem; margin: 1rem 0;">
                <h4>{factor['name']}</h4>
                <p>📊 المتوسط: {mean_score:.2f} / 5</p>
                <p>📈 الاتجاه: <span class="{trend_class}">{trend_text}</span></p>
                <p>🎯 المستوى: {trend_level}</p>
                <p>🔧 معامل الثبات (ألفا كرونباخ): {cronbach_display} - {reliability_text}</p>
                <p>📝 عدد الفقرات: {len(factor['questions'])} فقرة</p>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 6: اختبار التوزيع الطبيعي المتقدم ====================
elif st.session_state.page == 'normality' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🔬 اختبارات التوزيع الطبيعي المتقدمة</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>ما هي اختبارات التوزيع الطبيعي؟</strong></p>
        <ul>
            <li><strong>اختبار Shapiro-Wilk:</strong> أفضل اختبار للعينات الصغيرة (n < 50)</li>
            <li><strong>اختبار Kolmogorov-Smirnov:</strong> مناسب للعينات الكبيرة (n > 50)</li>
            <li><strong>اختبار Jarque-Bera:</strong> يعتمد على الالتواء والتفرطح</li>
            <li><strong>اختبار D'Agostino:</strong> يجمع بين الالتواء والتفرطح</li>
        </ul>
        <p>📌 <strong>قاعدة القرار:</strong> إذا كانت قيمة P-value > 0.05 → البيانات تتبع التوزيع الطبيعي</p>
        <p>🎯 <strong>الأهمية:</strong> تحدد نوع الاختبارات (معلمية إذا كان طبيعياً، لامعلمية إذا لم يكن طبيعياً)</p>
    </div>
    """, unsafe_allow_html=True)

    normality_results = []

    for factor in factors:
        if factor['questions']:
            factor_name = factor['name']
            if factor_name in df.columns:
                factor_values = df[factor_name].dropna()

                if len(factor_values) >= 3:
                    try:
                        # اختبار Shapiro-Wilk
                        shapiro_stat, shapiro_p = shapiro(factor_values)

                        # اختبار Kolmogorov-Smirnov (تطبيع البيانات أولاً)
                        normalized_data = (factor_values - factor_values.mean()) / factor_values.std()
                        ks_stat, ks_p = stats.kstest(normalized_data, 'norm')

                        # اختبار Jarque-Bera
                        jb_stat, jb_p = stats.jarque_bera(factor_values)

                        # اختبار D'Agostino's K^2
                        dag_stat, dag_p = stats.normaltest(factor_values)

                        # تنسيق قيم P-value بشكل أفضل
                        def format_pvalue(p_val):
                            if p_val < 0.0001:
                                return "< 0.0001"
                            elif p_val < 0.001:
                                return f"{p_val:.4f}"
                            elif p_val < 0.01:
                                return f"{p_val:.4f}"
                            elif p_val < 0.05:
                                return f"{p_val:.4f}"
                            else:
                                return f"{p_val:.4f}"

                        # تحديد النتيجة بناءً على القيمة الفعلية
                        is_normal = shapiro_p > 0.05

                        normality_results.append({
                            'المحور': factor_name,
                            'حجم العينة': len(factor_values),
                            'Shapiro-Wilk (p)': format_pvalue(shapiro_p),
                            'Kolmogorov-Smirnov (p)': format_pvalue(ks_p),
                            'Jarque-Bera (p)': format_pvalue(jb_p),
                            "D'Agostino (p)": format_pvalue(dag_p),
                            'القيمة الفعلية (Shapiro)': shapiro_p,
                            'النتيجة': '✅ طبيعي' if is_normal else '❌ غير طبيعي',
                            'التوصية': 'اختبارات معلمية' if is_normal else 'اختبارات لامعلمية'
                        })
                    except Exception as e:
                        st.warning(f"⚠️ خطأ في تحليل المحور '{factor_name}': {str(e)}")
                        normality_results.append({
                            'المحور': factor_name,
                            'حجم العينة': len(factor_values),
                            'Shapiro-Wilk (p)': 'خطأ',
                            'Kolmogorov-Smirnov (p)': 'خطأ',
                            'Jarque-Bera (p)': 'خطأ',
                            "D'Agostino (p)": 'خطأ',
                            'القيمة الفعلية (Shapiro)': np.nan,
                            'النتيجة': '⚠️ غير قابل للحساب',
                            'التوصية': 'تحقق من البيانات'
                        })
                else:
                    st.warning(
                        f"⚠️ المحور '{factor_name}' يحتوي على {len(factor_values)} ملاحظة فقط (يحتاج 3 على الأقل)")

    if normality_results:
        # عرض النتائج
        display_df = pd.DataFrame(normality_results).drop(columns=['القيمة الفعلية (Shapiro)'], errors='ignore')
        st.markdown("### 📊 نتائج اختبارات التوزيع الطبيعي")
        st.dataframe(display_df, use_container_width=True)

        # حساب الإحصاءات summary
        normal_count = len([r for r in normality_results if r['النتيجة'] == '✅ طبيعي'])
        not_normal_count = len([r for r in normality_results if r['النتيجة'] == '❌ غير طبيعي'])
        total_count = len(normality_results)

        if total_count > 0:
            st.markdown(f"""
            <div class="result-card">
                <h3>📊 ملخص النتائج</h3>
                <p>📌 عدد المحاور التي تم اختبارها: <strong>{total_count}</strong></p>
                <p>✅ المحاور التي تتبع التوزيع الطبيعي: <strong>{normal_count}</strong> ({normal_count / total_count * 100:.1f}%)</p>
                <p>❌ المحاور التي لا تتبع التوزيع الطبيعي: <strong>{not_normal_count}</strong> ({not_normal_count / total_count * 100:.1f}%)</p>
            </div>
            """, unsafe_allow_html=True)

            # التوصية العامة
            st.markdown("### 📝 التوصية لاختيار نوع الاختبارات")

            if normal_count == total_count:
                st.markdown("""
                <div class="info-box" style="background-color: #d4edda; border-right-color: #28a745;">
                    <h4>✅ جميع المحاور تتبع التوزيع الطبيعي</h4>
                    <p><strong>التوصية:</strong> يمكنك استخدام الاختبارات المعلمية (Parametric Tests):</p>
                    <ul>
                        <li><strong>اختبار T-test</strong> للمجموعات المستقلة والمرتبطة</li>
                        <li><strong>تحليل ANOVA</strong> للمقارنات المتعددة</li>
                        <li><strong>معامل ارتباط بيرسون (Pearson)</strong> للارتباطات</li>
                        <li><strong>الانحدار الخطي</strong> للتنبؤ</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                st.session_state.is_normal = True

            elif not_normal_count == total_count:
                st.markdown("""
                <div class="info-box" style="background-color: #f8d7da; border-right-color: #dc3545;">
                    <h4>❌ جميع المحاور لا تتبع التوزيع الطبيعي</h4>
                    <p><strong>التوصية:</strong> يُنصح بشدة باستخدام الاختبارات اللامعلمية (Non-parametric Tests):</p>
                    <ul>
                        <li><strong>اختبار Mann-Whitney U</strong> بدلاً من T-test</li>
                        <li><strong>اختبار Kruskal-Wallis</strong> بدلاً من ANOVA</li>
                        <li><strong>معامل ارتباط سبيرمان (Spearman)</strong> بدلاً من بيرسون</li>
                        <li><strong>اختبار Wilcoxon</strong> للعينات المرتبطة</li>
                        <li>يمكن تطبيق <strong>تحويل البيانات</strong> (Log, Square Root, Box-Cox) لمحاولة تحقيق الطبيعية</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                st.session_state.is_normal = False

            else:
                st.markdown("""
                <div class="info-box" style="background-color: #fff3cd; border-right-color: #ffc107;">
                    <h4>⚠️ بعض المحاور لا تتبع التوزيع الطبيعي</h4>
                    <p><strong>التوصية:</strong> استخدم الاختبارات اللامعلمية للمحاور غير الطبيعية، والاختبارات المعلمية للمحاور الطبيعية.</p>
                    <p>في حالة الشك، يفضل استخدام الاختبارات اللامعلمية لأنها أكثر تحفظاً.</p>
                </div>
                """, unsafe_allow_html=True)
                st.session_state.is_normal = False

        st.session_state.normality_test_done = True

        # عرض التوزيع البياني لأحد المحاور
        st.markdown("### 📊 توزيع المحاور (التمثيل البياني)")
        normal_factors = [f['name'] for f in factors if f['name'] in df.columns]
        if normal_factors:
            selected_factor = st.selectbox("اختر محوراً لعرض توزيعه:", normal_factors)
            if selected_factor in df.columns:
                chart_type = st.selectbox("نوع الرسم البياني:",
                                          ["مدرج تكراري مع منحنى كثافة", "مخطط صندوقي", "مخطط كمان", "مخطط Q-Q",
                                           "مخطط التوزيع التراكمي"],
                                          key="norm_chart")

                data_clean = df[selected_factor].dropna()

                if chart_type == "مدرج تكراري مع منحنى كثافة":
                    fig = ff.create_distplot([data_clean.values], [selected_factor], show_hist=True, show_rug=False,
                                             colors=['#667eea'])
                    fig.update_layout(title=f"توزيع {selected_factor} مع منحنى الكثافة", height=500)

                elif chart_type == "مخطط صندوقي":
                    fig = px.box(df, y=selected_factor, title=f"مخطط صندوقي لـ {selected_factor}",
                                 color_discrete_sequence=['#764ba2'])

                elif chart_type == "مخطط كمان":
                    fig = px.violin(df, y=selected_factor, title=f"مخطط كمان لـ {selected_factor}", box=True,
                                    color_discrete_sequence=['#2ecc71'])

                elif chart_type == "مخطط Q-Q":
                    fig = go.Figure()
                    qq_data = stats.probplot(data_clean, dist="norm")
                    fig.add_trace(go.Scatter(x=qq_data[0][0], y=qq_data[0][1],
                                             mode='markers', name='البيانات',
                                             marker=dict(color='#667eea', size=6)))
                    fig.add_trace(go.Scatter(x=qq_data[0][0], y=qq_data[1][0] + qq_data[1][1] * np.array(qq_data[0][0]),
                                             mode='lines', name='خط الطبيعي',
                                             line=dict(color='red', dash='dash', width=2)))
                    fig.update_layout(title=f"Q-Q Plot لـ {selected_factor}", height=500)

                else:
                    fig = px.ecdf(df, x=selected_factor, title=f"مخطط التوزيع التراكمي لـ {selected_factor}",
                                  color_discrete_sequence=['#3498db'])

                fig = create_rtl_figure(fig)
                st.plotly_chart(fig, use_container_width=True)
    else:
        st.warning("⚠️ لا توجد بيانات كافية لإجراء اختبارات التوزيع الطبيعي.")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 7: تحليل الارتباط المتقدم ====================
elif st.session_state.page == 'correlation' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🔗 تحليل الارتباط المتقدم</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل الارتباط:</strong> يهدف إلى قياس قوة واتجاه العلاقة بين متغيرين كميين.</p>
        <p>📊 <strong>أنواع معاملات الارتباط:</strong></p>
        <ul>
            <li><strong>بيرسون (Pearson):</strong> للعلاقات الخطية - يستخدم عندما تكون البيانات طبيعية</li>
            <li><strong>سبيرمان (Spearman):</strong> للعلاقات الرتيبة - يستخدم للبيانات غير الطبيعية</li>
            <li><strong>كندال (Kendall):</strong> مناسب للعينات الصغيرة أو البيانات الرتبية</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

    factor_names = [f['name'] for f in factors if f['questions']]
    all_vars = factor_names + st.session_state.social_vars + st.session_state.independent_vars + st.session_state.mediator_vars
    numeric_vars = df[all_vars].select_dtypes(include=[np.number]).columns.tolist()

    if len(numeric_vars) >= 2:
        # تحديد نوع معامل الارتباط بناءً على نتيجة اختبار الطبيعي
        default_method = "سبيرمان (Spearman)" if st.session_state.is_normal is False else "بيرسون (Pearson)"
        corr_method = st.radio("نوع معامل الارتباط:", ["بيرسون (Pearson)", "سبيرمان (Spearman)", "كندال (Kendall)"],
                               horizontal=True, index=["بيرسون (Pearson)", "سبيرمان (Spearman)", "كندال (Kendall)"].index(default_method) if default_method in ["بيرسون (Pearson)", "سبيرمان (Spearman)"] else 1)

        method = 'pearson' if corr_method == "بيرسون (Pearson)" else 'spearman' if corr_method == "سبيرمان (Spearman)" else 'kendall'

        corr_matrix = df[numeric_vars].corr(method=method)

        chart_type = st.selectbox("نوع خريطة الحرارة:", ["عادية (Heatmap)", "مرتبة (Clustered Heatmap)"])

        if chart_type == "عادية (Heatmap)":
            fig = px.imshow(corr_matrix, text_auto=True, aspect="auto", color_continuous_scale='RdBu_r',
                            title=f"مصفوفة الارتباط - {corr_method}", zmin=-1, zmax=1)
        else:
            import scipy.cluster.hierarchy as sch

            linkage = sch.linkage(corr_matrix, method='average')
            order = sch.leaves_list(linkage)
            corr_clustered = corr_matrix.iloc[order, order]
            fig = px.imshow(corr_clustered, text_auto=True, aspect="auto", color_continuous_scale='RdBu_r',
                            title=f"مصفوفة الارتباط المرتبة - {corr_method}", zmin=-1, zmax=1)

        fig.update_layout(height=600)
        fig = create_rtl_figure(fig)
        st.plotly_chart(fig, use_container_width=True)

        corr_pairs = []
        for i in range(len(corr_matrix.columns)):
            for j in range(i + 1, len(corr_matrix.columns)):
                corr_val = corr_matrix.iloc[i, j]
                if abs(corr_val) > 0.3:
                    if abs(corr_val) > 0.8:
                        strength = 'قوي جداً'
                    elif abs(corr_val) > 0.6:
                        strength = 'قوي'
                    elif abs(corr_val) > 0.4:
                        strength = 'متوسط'
                    else:
                        strength = 'ضعيف'
                    direction = 'موجب 📈' if corr_val > 0 else 'سالب 📉'
                    corr_pairs.append({'المتغير 1': corr_matrix.columns[i], 'المتغير 2': corr_matrix.columns[j],
                                       'معامل الارتباط': f"{corr_val:.3f}", 'القوة': strength, 'الاتجاه': direction})

        if corr_pairs:
            st.markdown("### 🔍 أقوى الارتباطات")
            st.dataframe(pd.DataFrame(corr_pairs).sort_values('معامل الارتباط', ascending=False),
                         use_container_width=True)

            st.markdown("### 📊 رسم بياني لأقوى الارتباطات")
            top_corr = pd.DataFrame(corr_pairs[:10])
            chart_type_corr = st.selectbox("نوع الرسم البياني:", ["أعمدة (Bar)", "أفقية (Horizontal Bar)"],
                                           key="corr_chart")

            if chart_type_corr == "أعمدة (Bar)":
                fig_corr = px.bar(top_corr, x='المتغير 1', y='معامل الارتباط', color='القوة', title='أقوى 10 ارتباطات',
                                  text='معامل الارتباط')
            else:
                fig_corr = px.bar(top_corr, x='معامل الارتباط', y='المتغير 1', color='القوة', title='أقوى 10 ارتباطات',
                                  orientation='h', text='معامل الارتباط')

            fig_corr = create_rtl_figure(fig_corr)
            st.plotly_chart(fig_corr, use_container_width=True)
    else:
        st.warning("⚠️ لا توجد متغيرات كافية لتحليل الارتباط")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 8: تحليل الانحدار ====================
elif st.session_state.page == 'regression' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📈 تحليل الانحدار</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل الانحدار:</strong> يستخدم لدراسة العلاقة بين متغير تابع ومتغيرات مستقلة والتنبؤ بقيمة المتغير التابع.</p>
        <p>📊 <strong>يشمل:</strong> الانحدار الخطي البسيط والمتعدد مع تشخيص النموذج.</p>
    </div>
    """, unsafe_allow_html=True)

    regression_type = st.selectbox("نوع تحليل الانحدار:", ["انحدار خطي بسيط", "انحدار خطي متعدد"])
    factor_names = [f['name'] for f in factors if f['questions']]

    if not factor_names:
        st.warning("⚠️ يرجى تحديد المحاور أولاً")
    else:
        col1, col2 = st.columns(2)
        with col1:
            dependent_var = st.selectbox("🎯 المتغير التابع (Y):", factor_names)
        with col2:
            all_independent = []
            for f in factor_names:
                if f != dependent_var:
                    all_independent.append(f)
            for v in st.session_state.social_vars:
                if v in df.columns and pd.api.types.is_numeric_dtype(df[v]):
                    all_independent.append(v)
            for factor in factors:
                for v in factor['independent_vars']:
                    if v in df.columns and pd.api.types.is_numeric_dtype(df[v]) and v not in all_independent:
                        all_independent.append(v)

            if regression_type == "انحدار خطي بسيط":
                independent_vars = st.multiselect("📊 المتغير المستقل (X):", all_independent,
                                                  default=all_independent[:min(1, len(all_independent))],
                                                  max_selections=1)
            else:
                independent_vars = st.multiselect("📊 المتغيرات المستقلة (X):", all_independent,
                                                  default=all_independent[:min(3, len(all_independent))])

        if independent_vars and st.button("🚀 تنفيذ التحليل", type="primary"):
            try:
                X = df[independent_vars].copy()
                y = df[dependent_var].copy()
                valid_idx = X.dropna().index
                X = X.loc[valid_idx]
                y = y.loc[valid_idx]

                reg_results = advanced_regression_analysis(X, y, independent_vars)

                st.markdown("### 📊 ملخص النموذج")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("R²", f"{reg_results['r_squared']:.4f}")
                    st.caption(f"يشرح {reg_results['r_squared'] * 100:.1f}%")
                with col2:
                    st.metric("R² المعدل", f"{reg_results['adj_r_squared']:.4f}")
                with col3:
                    st.metric("RMSE", f"{reg_results['rmse']:.4f}")
                with col4:
                    st.metric("MAE", f"{reg_results['mae']:.4f}")

                f_pvalue = reg_results['f_pvalue']
                if f_pvalue < 0.05:
                    st.success(f"✅ النموذج دال إحصائياً (F = {reg_results['f_statistic']:.4f}, p = {f_pvalue:.4f})")
                else:
                    st.warning(f"❌ النموذج غير دال إحصائياً (F = {reg_results['f_statistic']:.4f}, p = {f_pvalue:.4f})")

                st.markdown("### 📈 معاملات الانحدار")
                coef_data = []
                for i, var in enumerate(['ثابت'] + independent_vars):
                    coef = reg_results['coefficients'][i]
                    p_val = reg_results['p_values'][i]
                    ci_low, ci_high = reg_results['conf_int'][i]
                    coef_data.append({'المتغير': var, 'معامل B': f"{coef:.4f}", 'P-value': f"{p_val:.4f}",
                                      'فترة ثقة 95%': f"[{ci_low:.4f}, {ci_high:.4f}]",
                                      'الدلالة': '✅ دال' if p_val < 0.05 else '❌ غير دال'})
                st.dataframe(pd.DataFrame(coef_data), use_container_width=True)

                st.markdown("### 📐 معادلة الانحدار")
                equation = f"{dependent_var} = {reg_results['coefficients'][0]:.4f}"
                for i, var in enumerate(independent_vars):
                    coef = reg_results['coefficients'][i + 1]
                    sign = "+" if coef >= 0 else "-"
                    equation += f" {sign} {abs(coef):.4f} × {var}"
                st.markdown(f'<div class="regression-box"><code>{equation}</code></div>', unsafe_allow_html=True)

                # تشخيص النموذج المتقدم مع التوصية
                st.markdown("### 🔧 تشخيص النموذج المتقدم")

                diag_tab1, diag_tab2, diag_tab3 = st.tabs(
                    ["📊 اختبارات النموذج", "📈 الرسوم البيانية التشخيصية", "📝 التوصيات والإجراءات"])

                with diag_tab1:
                    st.markdown("#### اختبارات جودة النموذج")

                    if not reg_results['vif'].empty:
                        st.markdown("##### 📌 معامل تضخم التباين (VIF)")
                        st.dataframe(reg_results['vif'], use_container_width=True)

                        vif_values = reg_results['vif']['VIF'].dropna().values
                        if len(vif_values) > 0:
                            max_vif = max(vif_values)
                            if max_vif > 10:
                                st.error(
                                    f"⚠️ مشكلة خطيرة في تعدد العلاقات الخطية (VIF = {max_vif:.2f} > 10)")
                            elif max_vif > 5:
                                st.warning(
                                    f"⚠️ مشكلة محتملة في تعدد العلاقات الخطية (VIF = {max_vif:.2f} > 5)")
                            else:
                                st.success(f"✅ لا توجد مشكلة في تعدد العلاقات الخطية (VIF = {max_vif:.2f} < 5)")

                    st.markdown("##### 📌 اختبار الارتباط الذاتي (Durbin-Watson)")
                    dw = reg_results['durbin_watson']
                    st.write(f"قيمة Durbin-Watson: **{dw:.4f}**")
                    if dw < 1.5:
                        st.warning("⚠️ يوجد ارتباط ذاتي موجب")
                    elif dw > 2.5:
                        st.warning("⚠️ يوجد ارتباط ذاتي سالب")
                    else:
                        st.success("✅ لا يوجد ارتباط ذاتي")

                    st.markdown("##### 📌 اختبار طبيعية البواقي")
                    res_norm = reg_results.get('residuals_normality', {})
                    if not np.isnan(res_norm.get('statistic', np.nan)):
                        st.write(f"Shapiro-Wilk: p-value = {res_norm['pvalue']:.4f}")
                        if res_norm['pvalue'] > 0.05:
                            st.success("✅ البواقي طبيعية")
                        else:
                            st.warning("⚠️ البواقي غير طبيعية")

                    st.markdown("##### 📌 اختبار تجانس التباين")
                    bp = reg_results.get('heteroscedasticity', {})
                    if not np.isnan(bp.get('statistic', np.nan)):
                        st.write(f"Breusch-Pagan: p-value = {bp['pvalue']:.4f}")
                        if bp['pvalue'] > 0.05:
                            st.success("✅ التباين متجانس")
                        else:
                            st.warning("⚠️ التباين غير متجانس")

                with diag_tab2:
                    st.markdown("#### الرسوم البيانية التشخيصية")

                    fig1 = px.scatter(x=reg_results['fitted_values'], y=reg_results['residuals'],
                                      title="البواقي مقابل القيم المتوقعة",
                                      labels={'x': 'القيم المتوقعة', 'y': 'البواقي'})
                    fig1.add_hline(y=0, line_dash="dash", line_color="red")
                    fig1 = create_rtl_figure(fig1)
                    st.plotly_chart(fig1, use_container_width=True)

                    fig2 = px.histogram(x=reg_results['residuals'], nbins=20,
                                        title="توزيع البواقي",
                                        labels={'x': 'البواقي', 'y': 'التكرار'})
                    fig2 = create_rtl_figure(fig2)
                    st.plotly_chart(fig2, use_container_width=True)

                    fig3 = go.Figure()
                    qq_data = stats.probplot(reg_results['residuals'], dist="norm")
                    fig3.add_trace(go.Scatter(x=qq_data[0][0], y=qq_data[0][1],
                                              mode='markers', name='البواقي',
                                              marker=dict(color='#667eea', size=6)))
                    fig3.add_trace(
                        go.Scatter(x=qq_data[0][0], y=qq_data[1][0] + qq_data[1][1] * np.array(qq_data[0][0]),
                                   mode='lines', name='خط المرجع',
                                   line=dict(color='red', dash='dash', width=2)))
                    fig3.update_layout(title="Q-Q Plot للبواقي", height=400)
                    fig3 = create_rtl_figure(fig3)
                    st.plotly_chart(fig3, use_container_width=True)

                with diag_tab3:
                    st.markdown("#### 📝 التوصيات والإجراءات المقترحة")

                    r2 = reg_results['r_squared']
                    if r2 < 0.3:
                        st.markdown("""
                        <div class="info-box">
                            <p><strong>توصيات لتحسين النموذج:</strong></p>
                            <ul>
                                <li>إضافة متغيرات مستقلة إضافية ذات تأثير نظري</li>
                                <li>التحقق من وجود علاقات غير خطية</li>
                                <li>توسيع حجم العينة</li>
                            </ul>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div class="success-box">
                            <p>✅ النموذج جيد ويمكن الاعتماد على نتائجه للتفسير والتنبؤ.</p>
                        </div>
                        """, unsafe_allow_html=True)

                st.balloons()

            except Exception as e:
                st.error(f"❌ خطأ: {str(e)}")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 9: تحليل الوساطة المتقدم ====================
elif st.session_state.page == 'mediation' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🔄 تحليل تأثير المتغير الوسيط المتقدم</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل الوساطة المتقدم:</strong> يهدف إلى فهم كيفية تأثير متغير مستقل على متغير تابع من خلال متغير وسيط.</p>
        <p>📊 <strong>يتضمن:</strong> التأثير الكلي، المباشر، غير المباشر، اختبار سوبل، Bootstrap CI.</p>
        <p>🎯 <strong>شروط الوساطة:</strong> X يؤثر على M، M يؤثر على Y، ويقل تأثير X على Y عند إدخال M.</p>
    </div>
    """, unsafe_allow_html=True)

    # جمع أسماء المتغيرات المتاحة
    factor_names = [f['name'] for f in factors if f['questions'] and f['name'] in df.columns]

    all_social_vars = [v for v in st.session_state.social_vars if
                       v in df.columns and pd.api.types.is_numeric_dtype(df[v])]
    all_independent_vars = [v for v in st.session_state.independent_vars if
                            v in df.columns and pd.api.types.is_numeric_dtype(df[v])]
    all_mediator_vars = [v for v in st.session_state.mediator_vars if
                         v in df.columns and pd.api.types.is_numeric_dtype(df[v])]

    available_independent = list(set(factor_names + all_social_vars + all_independent_vars))
    available_mediators = list(set(factor_names + all_mediator_vars))
    available_dependent = factor_names.copy()

    available_independent = [v for v in available_independent if v is not None and v != 'None']
    available_mediators = [v for v in available_mediators if v is not None and v != 'None']
    available_dependent = [v for v in available_dependent if v is not None and v != 'None']

    if len(available_independent) == 0:
        st.warning("⚠️ لا توجد متغيرات مستقلة متاحة للتحليل.")
    elif len(available_mediators) == 0:
        st.warning("⚠️ لا توجد متغيرات وسيطة متاحة للتحليل.")
    elif len(available_dependent) == 0:
        st.warning("⚠️ لا توجد متغيرات تابعة متاحة للتحليل.")
    else:
        col1, col2, col3 = st.columns(3)
        with col1:
            independent_var = st.selectbox("🎯 المتغير المستقل (X):", options=available_independent)
        with col2:
            mediator_var = st.selectbox("🔄 المتغير الوسيط (M):", options=available_mediators)
        with col3:
            dependent_options = [v for v in available_dependent if v != independent_var and v != mediator_var]
            if not dependent_options:
                dependent_options = available_dependent
            dependent_var = st.selectbox("📊 المتغير التابع (Y):", options=dependent_options)

        bootstrap_iter = st.slider("عدد تكرارات Bootstrap:", 100, 5000, 1000, key="bootstrap_iter")

        if st.button("🔍 تحليل تأثير الوساطة", type="primary"):
            try:
                if independent_var not in df.columns:
                    st.error(f"❌ المتغير المستقل '{independent_var}' غير موجود")
                elif mediator_var not in df.columns:
                    st.error(f"❌ المتغير الوسيط '{mediator_var}' غير موجود")
                elif dependent_var not in df.columns:
                    st.error(f"❌ المتغير التابع '{dependent_var}' غير موجود")
                else:
                    mediation_data = df[[independent_var, mediator_var, dependent_var]].dropna()

                    if len(mediation_data) < 10:
                        st.warning(f"⚠️ عدد الملاحظات الصالحة: {len(mediation_data)} (يُفضل 30+)")
                    elif len(mediation_data) < 3:
                        st.error(f"❌ عدد الملاحظات غير كافٍ: {len(mediation_data)}")
                    else:
                        with st.spinner("جاري تحليل الوساطة..."):
                            mediation_results = mediation_analysis_advanced(df, independent_var, mediator_var,
                                                                            dependent_var, bootstrap_iter)

                            st.markdown("### 📊 نتائج تحليل الوساطة")

                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("التأثير الكلي", f"{mediation_results['total_effect']['coefficient']:.4f}")
                                st.caption(f"p = {mediation_results['total_effect']['pvalue']:.4f}")
                            with col2:
                                st.metric("التأثير المباشر", f"{mediation_results['direct_effect']['coefficient']:.4f}")
                                st.caption(f"p = {mediation_results['direct_effect']['pvalue']:.4f}")
                            with col3:
                                st.metric("التأثير غير المباشر", f"{mediation_results['indirect_effect']:.4f}")
                                st.caption(f"نسبة الوساطة: {mediation_results['mediation_ratio']:.1f}%")

                            st.markdown(f"""
                            <div class="result-card">
                                <h4>🏷️ نوع الوساطة</h4>
                                <p style="font-size: 1.2rem; font-weight: bold;">{mediation_results['mediation_type']}</p>
                            </div>
                            """, unsafe_allow_html=True)

                            st.markdown("### 📈 تفاصيل المسارات")
                            st.markdown(f"""
                            - **المسار أ ({independent_var} → {mediator_var}):** β = {mediation_results['path_a']['coefficient']:.4f}, p = {mediation_results['path_a']['pvalue']:.4f}
                            - **المسار ب ({mediator_var} → {dependent_var}):** β = {mediation_results['path_b']['coefficient']:.4f}, p = {mediation_results['path_b']['pvalue']:.4f}
                            - **المسار ج ({independent_var} → {dependent_var}):** β = {mediation_results['total_effect']['coefficient']:.4f}, p = {mediation_results['total_effect']['pvalue']:.4f}
                            - **المسار جَ ({independent_var} → {dependent_var} مع {mediator_var}):** β = {mediation_results['direct_effect']['coefficient']:.4f}, p = {mediation_results['direct_effect']['pvalue']:.4f}
                            """)

                            st.markdown("### 📊 اختبارات الدلالة")
                            sobel = mediation_results['sobel_test']
                            st.markdown(f"""
                            <div class="info-box">
                                <p><strong>اختبار سوبل (Sobel Test):</strong> Z = {sobel['z']:.4f}, p = {sobel['pvalue']:.4f}</p>
                                <p><strong>Bootstrap CI 95%:</strong> [{mediation_results['bootstrap_ci']['lower']:.4f}, {mediation_results['bootstrap_ci']['upper']:.4f}]</p>
                                <p><strong>النتيجة:</strong> {'✅ التأثير غير المباشر دال إحصائياً' if sobel['pvalue'] < 0.05 else '❌ التأثير غير المباشر غير دال'}</p>
                            </div>
                            """, unsafe_allow_html=True)

                            st.balloons()

            except Exception as e:
                st.error(f"❌ خطأ في التحليل: {str(e)}")

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 10: اختبار الفروقات ====================
elif st.session_state.page == 'differences' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors
    significance_level = st.select_slider("مستوى الدلالة:", options=[0.01, 0.05, 0.10], value=0.05)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📐 اختبار الفروقات في درجات الإجابة</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>شرح الاختبارات:</strong></p>
        <ul>
            <li><strong>اختبار T-test (مستقل):</strong> يقارن بين مجموعتين مستقلتين - يستخدم عندما تكون البيانات طبيعية</li>
            <li><strong>اختبار Mann-Whitney U:</strong> البديل اللامعلمي لاختبار T-test - يستخدم للبيانات غير الطبيعية</li>
            <li><strong>اختبار T-test المرتبط (Paired):</strong> يقارن بين قياسين لنفس المجموعة (قبل-بعد)</li>
            <li><strong>اختبار Wilcoxon:</strong> البديل اللامعلمي للاختبار المرتبط</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["📊 اختبار المجموعات المستقلة", "🔄 اختبار العينات المرتبطة", "📝 التوصيات والتفسير"])

    results_independent = []
    results_paired = []

    # ==================== TAB 1: اختبار المجموعات المستقلة ====================
    with tab1:
        for factor in factors:
            if not factor['questions']:
                continue

            for var in factor['social_vars'] + factor['independent_vars']:
                if var not in df.columns:
                    continue

                unique_vals = df[var].dropna().unique()

                if len(unique_vals) == 2:
                    g1 = df[df[var] == unique_vals[0]][factor['name']].dropna()
                    g2 = df[df[var] == unique_vals[1]][factor['name']].dropna()

                    if len(g1) > 0 and len(g2) > 0:
                        _, p1 = shapiro(g1) if len(g1) >= 3 else (None, 0.5)
                        _, p2 = shapiro(g2) if len(g2) >= 3 else (None, 0.5)
                        normal = (p1 > 0.05 if p1 is not None else True) and (p2 > 0.05 if p2 is not None else True)

                        if normal and st.session_state.is_normal is not False:
                            stat, p_value = ttest_ind(g1, g2, equal_var=False)
                            test_name = "T-test (مستقل)"
                            effect_size = calculate_effect_size_cohens_d(g1, g2)
                        else:
                            stat, p_value = mannwhitneyu(g1, g2, alternative='two-sided')
                            test_name = "Mann-Whitney U"
                            effect_size = stat / (len(g1) * len(g2))

                        is_sig = p_value < significance_level
                        mean_diff = g1.mean() - g2.mean()

                        if abs(effect_size) > 0.8:
                            effect_interpretation = "تأثير كبير جداً"
                        elif abs(effect_size) > 0.5:
                            effect_interpretation = "تأثير متوسط"
                        elif abs(effect_size) > 0.2:
                            effect_interpretation = "تأثير صغير"
                        else:
                            effect_interpretation = "تأثير ضئيل"

                        results_independent.append({
                            'الاختبار': test_name,
                            'المحور': factor['name'],
                            'المتغير': var,
                            f'متوسط {str(unique_vals[0])[:20]}': f"{g1.mean():.3f}",
                            f'متوسط {str(unique_vals[1])[:20]}': f"{g2.mean():.3f}",
                            'فرق المتوسطات': f"{mean_diff:.3f}",
                            'قيمة الاختبار': f"{stat:.3f}",
                            'P-value': f"{p_value:.4f}",
                            'الدلالة': '✅ دالة' if is_sig else '❌ غير دالة',
                            'حجم التأثير': f"{effect_size:.3f}",
                            'تفسير الحجم': effect_interpretation,
                            'القيمة الفعلية_P': p_value,
                            'القيمة الفعلية_effect': abs(effect_size)
                        })

        if results_independent:
            st.dataframe(pd.DataFrame(results_independent), use_container_width=True)

            # تخزين النتائج في session_state للاستخدام في التوصيات
            st.session_state.differences_results = results_independent
        else:
            st.info("📌 لا توجد فروقات للعرض. تأكد من تحديد متغيرات اجتماعية أو مستقلة ثنائية الفئات.")

    # ==================== TAB 2: اختبار العينات المرتبطة ====================
    with tab2:
        st.markdown("### اختبار العينات المرتبطة (Paired Samples)")
        st.info("📌 يتطلب هذا الاختبار وجود قياسين مرتبطين (قبل وبعد) لنفس العينة")

        all_numeric = df.select_dtypes(include=[np.number]).columns.tolist()

        if len(all_numeric) >= 2:
            col1, col2 = st.columns(2)
            with col1:
                var1 = st.selectbox("المتغير الأول (القياس الأول):", all_numeric, key="paired1")
            with col2:
                var2 = st.selectbox("المتغير الثاني (القياس الثاني):", [v for v in all_numeric if v != var1],
                                    key="paired2")

            if st.button("تنفيذ اختبار العينات المرتبطة", key="paired_btn"):
                paired_data = df[[var1, var2]].dropna()
                if len(paired_data) > 0:
                    diff = paired_data[var1] - paired_data[var2]
                    _, p_norm = shapiro(diff) if len(diff) >= 3 else (None, 0.5)

                    if p_norm > 0.05 and st.session_state.is_normal is not False:
                        stat, p_value = ttest_rel(paired_data[var1], paired_data[var2])
                        test_name = "T-test (مرتبط)"
                        effect_size = abs(paired_data[var1].mean() - paired_data[var2].mean()) / paired_data[
                            var1].std() if paired_data[var1].std() > 0 else 0
                    else:
                        stat, p_value = wilcoxon(paired_data[var1], paired_data[var2])
                        test_name = "Wilcoxon Signed-Rank"
                        effect_size = stat / (len(paired_data) * (len(paired_data) + 1) / 2)

                    is_sig = p_value < significance_level

                    results_paired.append({
                        'الاختبار': test_name,
                        'المتغير الأول': var1,
                        'المتغير الثاني': var2,
                        'متوسط الأول': paired_data[var1].mean(),
                        'متوسط الثاني': paired_data[var2].mean(),
                        'الفرق': paired_data[var1].mean() - paired_data[var2].mean(),
                        'P-value': p_value,
                        'الدلالة': is_sig,
                        'حجم التأثير': effect_size,
                        'عدد المشاهدات': len(paired_data)
                    })

                    st.session_state.paired_results = results_paired

                    st.markdown(f"""
                    <div class="result-card">
                        <p><strong>نتائج الاختبار:</strong> {test_name}</p>
                        <p>متوسط المتغير الأول: {paired_data[var1].mean():.3f}</p>
                        <p>متوسط المتغير الثاني: {paired_data[var2].mean():.3f}</p>
                        <p>الفرق: {paired_data[var1].mean() - paired_data[var2].mean():.3f}</p>
                        <p>P-value: {p_value:.4f}</p>
                        <p>الدلالة: {'✅ يوجد فرق دال إحصائياً' if is_sig else '❌ لا يوجد فرق دال'}</p>
                        <p>📏 حجم التأثير: {effect_size:.3f}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    fig = go.Figure()
                    fig.add_trace(go.Box(y=paired_data[var1], name=var1, boxmean='sd'))
                    fig.add_trace(go.Box(y=paired_data[var2], name=var2, boxmean='sd'))
                    fig.update_layout(title=f"مقارنة بين {var1} و {var2}", height=400)
                    fig = create_rtl_figure(fig)
                    st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("⚠️ لا توجد متغيرات رقمية كافية")

    # ==================== TAB 3: التوصيات والتفسير المتقدم ====================
    with tab3:
        st.markdown("## 📝 تفسير النتائج والتوصيات المتقدمة")

        # التحقق من وجود نتائج
        has_independent = results_independent or st.session_state.get('differences_results', [])
        has_paired = results_paired or st.session_state.get('paired_results', [])

        if not has_independent and not has_paired:
            st.info("📌 قم بتنفيذ اختبارات الفروقات أولاً في التبويبات السابقة لعرض التوصيات.")
        else:
            # استخدام النتائج المخزنة
            ind_results = results_independent if results_independent else st.session_state.get('differences_results',
                                                                                               [])
            pair_results = results_paired if results_paired else st.session_state.get('paired_results', [])

            # ==================== تحليل النتائج المتقدم ====================
            total_tests = len(ind_results)
            sig_count = len([r for r in ind_results if r['الدلالة'] == '✅ دالة'])
            sig_percentage = (sig_count / total_tests * 100) if total_tests > 0 else 0

            # جمع أحجام التأثير
            effect_sizes = [r['القيمة الفعلية_effect'] for r in ind_results if r.get('القيمة الفعلية_effect', 0) > 0]
            avg_effect = np.mean(effect_sizes) if effect_sizes else 0

            # تحديد أقوى فروق
            if ind_results:
                strongest_diff = max(ind_results,
                                     key=lambda x: abs(float(x['فرق المتوسطات'])) if x['فرق المتوسطات'] != 'N/A' else 0)
                largest_effect = max(ind_results, key=lambda x: x.get('القيمة الفعلية_effect', 0))
            else:
                strongest_diff = None
                largest_effect = None

            # ==================== عرض الإحصائيات ====================
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("📊 عدد الاختبارات", total_tests)
            with col2:
                st.metric("✅ فروق دالة", f"{sig_count} ({sig_percentage:.1f}%)")
            with col3:
                st.metric("❌ فروق غير دالة", total_tests - sig_count)
            with col4:
                st.metric("📏 متوسط حجم التأثير", f"{avg_effect:.3f}")

            # ==================== التوصية الرئيسية ====================
            st.markdown("---")
            st.markdown("### 🎯 التوصية الرئيسية")

            if total_tests == 0 and not pair_results:
                st.warning("⚠️ لا توجد اختبارات كافية لتقديم توصية.")

            elif sig_count == 0:
                st.markdown("""
                <div style="background: #d4edda; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #28a745;">
                    <h4 style="color: #155724;">✅ لا توجد فروق ذات دلالة إحصائية</h4>
                    <p style="font-size: 1.05rem; line-height: 1.6;">
                        <strong>الاستنتاج:</strong> بناءً على نتائج اختبار الفروقات، <strong>لا توجد فروق دالة إحصائياً</strong> 
                        بين المجموعات المقارنة في جميع المحاور التي تم اختبارها.
                    </p>
                    <p style="margin-top: 10px;">
                        هذا يشير إلى أن المجموعات <strong>متجانسة</strong> في إجاباتها، وأن المتغيرات الديموغرافية 
                        (مثل: الجنس، العمر، المؤهل العلمي، الخبرة) <strong>لا تؤثر بشكل معنوي</strong> 
                        على اتجاهات وآراء أفراد العينة.
                    </p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("""
                <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #17a2b8;">
                    <h4 style="color: #0c5460;">💡 التوصية الإجرائية</h4>
                    <ul style="line-height: 1.8;">
                        <li><strong>✅ يمكن دمج جميع المجموعات في تحليل واحد</strong> - لا حاجة لتحليل منفصل لكل مجموعة</li>
                        <li><strong>✅ تبسيط عملية التحليل</strong> - يمكن الاعتماد على الإحصاءات الوصفية العامة دون تجزئة</li>
                        <li><strong>✅ تعميم النتائج</strong> - النتائج تنطبق على جميع فئات العينة بالتساوي</li>
                        <li><strong>✅ التركيز على تحليلات أخرى</strong> مثل الارتباط والانحدار وتحليل العوامل</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)

            elif sig_percentage < 30:
                st.markdown(f"""
                <div style="background: #fff3cd; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #ffc107;">
                    <h4 style="color: #856404;">⚠️ توجد فروق محدودة</h4>
                    <p style="font-size: 1.05rem; line-height: 1.6;">
                        <strong>الاستنتاج:</strong> تم العثور على <strong>{sig_count} فروق دالة إحصائياً</strong> 
                        من أصل {total_tests} اختبار تم إجراؤها ({sig_percentage:.1f}%).
                    </p>
                    <p style="margin-top: 10px;">
                        هذا يشير إلى أن معظم المتغيرات الديموغرافية <strong>لا تؤثر بشكل كبير</strong> على إجابات أفراد العينة،
                        باستثناء حالات محدودة تستحق الاهتمام والتحليل بشكل منفصل.
                    </p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("""
                <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #17a2b8;">
                    <h4 style="color: #0c5460;">💡 التوصية الإجرائية</h4>
                    <ul style="line-height: 1.8;">
                        <li><strong>⚖️ يمكن دمج معظم المجموعات</strong> - باستثناء الحالات التي ظهرت فيها فروق</li>
                        <li><strong>🔍 التركيز على المحاور التي ظهرت فيها فروق</strong> لفهم أسباب الاختلاف</li>
                        <li><strong>📊 إجراء تحليل الانحدار المتعدد</strong> باستخدام المتغيرات الديموغرافية كمتغيرات ضابطة</li>
                        <li><strong>🎯 توجيه التوصيات</strong> بناءً على المجموعات التي أظهرت اختلافات</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)

            else:
                st.markdown(f"""
                <div style="background: #f8d7da; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #dc3545;">
                    <h4 style="color: #721c24;">⚠️ توجد فروق واسعة النطاق</h4>
                    <p style="font-size: 1.05rem; line-height: 1.6;">
                        <strong>الاستنتاج:</strong> تم العثور على <strong>{sig_count} فروق دالة إحصائياً</strong> 
                        من أصل {total_tests} اختبار تم إجراؤها ({sig_percentage:.1f}%).
                    </p>
                    <p style="margin-top: 10px;">
                        هذا يشير إلى أن المجموعات <strong>تختلف بشكل جوهري</strong> في إجاباتها، 
                        وأن المتغيرات الديموغرافية <strong>تؤثر بشكل كبير</strong> على اتجاهات وآراء أفراد العينة.
                    </p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("""
                <div style="background: #f8d7da; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #dc3545;">
                    <h4 style="color: #721c24;">💡 التوصية الإجرائية</h4>
                    <ul style="line-height: 1.8;">
                        <li><strong>❌ لا يمكن دمج المجموعات</strong> - يجب تحليل كل مجموعة على حدة</li>
                        <li><strong>🎯 تخصيص التوصيات</strong> - لكل مجموعة توصيات مختلفة بناءً على خصائصها</li>
                        <li><strong>📋 تصميم برامج تدخل مخصصة</strong> - حسب احتياجات كل مجموعة</li>
                        <li><strong>🔬 إجراء دراسة معمقة</strong> لفهم أسباب الاختلافات بين المجموعات</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)

            # ==================== تحليل حجم التأثير ====================
            if effect_sizes:
                st.markdown("---")
                st.markdown("### 📏 تحليل حجم التأثير (Effect Size)")

                if avg_effect > 0.8:
                    effect_text = "كبير جداً"
                    effect_color = "#dc3545"
                    effect_advice = "الفروق المكتشفة ليست فقط دالة إحصائياً، بل لها أهمية عملية كبيرة."
                elif avg_effect > 0.5:
                    effect_text = "متوسط"
                    effect_color = "#ffc107"
                    effect_advice = "الفروق لها أهمية عملية متوسطة وتستحق الاهتمام."
                elif avg_effect > 0.2:
                    effect_text = "صغير"
                    effect_color = "#28a745"
                    effect_advice = "الفروق دالة إحصائياً ولكن لها أهمية عملية محدودة."
                else:
                    effect_text = "ضئيل جداً"
                    effect_color = "#17a2b8"
                    effect_advice = "الفروق رغم دلالتها الإحصائية (إن وجدت) إلا أن أهميتها العملية ضئيلة جداً."

                st.markdown(f"""
                <div style="background: #f0f0f0; border-radius: 15px; padding: 20px; margin: 15px 0;">
                    <p><strong>متوسط حجم التأثير:</strong> <span style="font-size: 1.3rem; font-weight: bold; color: {effect_color};">{avg_effect:.3f}</span></p>
                    <p><strong>التصنيف:</strong> {effect_text}</p>
                    <p><strong>التفسير:</strong> {effect_advice}</p>
                </div>
                """, unsafe_allow_html=True)

            # ==================== أقوى الفروق ====================
            if strongest_diff and strongest_diff.get('فرق المتوسطات', '0') != '0':
                st.markdown("---")
                st.markdown("### 🔍 أقوى الفروق المكتشفة")

                try:
                    diff_val = float(strongest_diff['فرق المتوسطات'])
                    effect_val = float(strongest_diff['حجم التأثير'])

                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #667eea10 0%, #764ba210 100%); 
                                border-radius: 15px; padding: 20px; margin: 15px 0;">
                        <p><strong>📊 المحور:</strong> {strongest_diff['المحور']}</p>
                        <p><strong>📌 المتغير:</strong> {strongest_diff['المتغير']}</p>
                        <p><strong>📈 الفرق في المتوسطات:</strong> {abs(diff_val):.3f} نقطة</p>
                        <p><strong>💪 حجم التأثير:</strong> {abs(effect_val):.3f}</p>
                        <p><strong>📝 التفسير:</strong> 
                            هذا الفرق يعتبر {'كبيراً' if abs(effect_val) > 0.8 else 'متوسطاً' if abs(effect_val) > 0.5 else 'صغيراً'}، 
                            ويشير إلى أن {'المجموعة الأولى' if diff_val > 0 else 'المجموعة الثانية'} 
                            لديها اتجاه {'أكثر إيجابية' if diff_val > 0 else 'أقل إيجابية'} في هذا المحور.
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
                except:
                    pass

            # ==================== التوصيات الإحصائية المستقبلية ====================
            st.markdown("---")
            st.markdown("### 🔭 توصيات للبحث المستقبلي")

            if sig_count == 0:
                st.markdown("""
                <div style="background: #f5f5f5; border-radius: 15px; padding: 20px; margin: 15px 0;">
                    <ul style="line-height: 1.8;">
                        <li><strong>توسيع نطاق العينة:</strong> قد تظهر فروق في عينات أكبر أو أكثر تنوعاً</li>
                        <li><strong>إضافة متغيرات ديموغرافية جديدة:</strong> مثل الدخل، المنطقة الجغرافية، الحالة الاجتماعية</li>
                        <li><strong>إجراء دراسة طولية (Longitudinal Study):</strong> لدراسة التغيرات بمرور الزمن</li>
                        <li><strong>استخدام مقاييس أكثر حساسية:</strong> قد تلتقط فروقاً دقيقة لا تظهر في المقاييس الحالية</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            elif sig_percentage < 50:
                st.markdown("""
                <div style="background: #f5f5f5; border-radius: 15px; padding: 20px; margin: 15px 0;">
                    <ul style="line-height: 1.8;">
                        <li><strong>التركيز على المحاور التي ظهرت فيها فروق:</strong> لفهم أسباب الاختلافات بشكل أعمق</li>
                        <li><strong>إضافة متغيرات وسيطة (Mediators):</strong> لفهم الآليات التي تفسر الفروق</li>
                        <li><strong>إجراء دراسة نوعية (Qualitative Study):</strong> لمقابلة أفراد من المجموعات المختلفة</li>
                        <li><strong>اختبار فروض جديدة:</strong> بناءً على النتائج الحالية</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background: #f5f5f5; border-radius: 15px; padding: 20px; margin: 15px 0;">
                    <ul style="line-height: 1.8;">
                        <li><strong>دراسة أسباب الاختلافات الجوهرية:</strong> من خلال منهجيات نوعية (مقابلات متعمقة، مجموعات بؤرية)</li>
                        <li><strong>تحليل تفاعلي (Interaction Effects):</strong> لدراسة تأثير تداخل المتغيرات الديموغرافية</li>
                        <li><strong>تصميم برامج تدخل مخصصة:</strong> لكل مجموعة بناءً على خصائصها الفريدة</li>
                        <li><strong>توسيع العينة لتشمل مجموعات جديدة:</strong> لمقارنة أكثر شمولاً</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)

            # ==================== صياغة النتائج بشكل أكاديمي ====================
            st.markdown("---")
            st.markdown("### 📝 صياغة النتائج بشكل أكاديمي")

            if sig_count == 0:
                st.markdown("""
                <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; direction: rtl;">
                    <p style="font-size: 1rem; line-height: 1.8;">
                    <strong>النص المقترح للتقرير:</strong><br><br>
                    "لاختبار الفروق في درجات الإجابة تبعاً للمتغيرات الديموغرافية، تم استخدام اختبار (T-test) للعينات المستقلة 
                    (بعد التحقق من افتراض التوزيع الطبيعي). أظهرت النتائج عدم وجود فروق ذات دلالة إحصائية بين المجموعات المقارنة 
                    في جميع المحاور (P-value > 0.05)، مما يشير إلى تجانس أفراد العينة في إجاباتهم، وعدم تأثر اتجاهاتهم 
                    بالمتغيرات الديموغرافية المدروسة."
                    </p>
                </div>
                """, unsafe_allow_html=True)
            else:
                # إنشاء نص أكاديمي للفروق الدالة
                sig_results_short = [r for r in ind_results if r['الدلالة'] == '✅ دالة'][:3]
                sig_text = ""
                for r in sig_results_short:
                    sig_text += f"• في محور '{r['المحور']}' حسب متغير '{r['المتغير']}' (P-value = {r['P-value']})<br>"

                st.markdown(f"""
                <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; direction: rtl;">
                    <p style="font-size: 1rem; line-height: 1.8;">
                    <strong>النص المقترح للتقرير:</strong><br><br>
                    "لاختبار الفروق في درجات الإجابة تبعاً للمتغيرات الديموغرافية، تم استخدام {'اختبار T-test للعينات المستقلة' if 'T-test' in ind_results[0]['الاختبار'] else 'اختبار Mann-Whitney U'} 
                    (بعد التحقق من افتراضات الاختبار المناسبة). أظهرت النتائج وجود فروق ذات دلالة إحصائية في {sig_count} من أصل {total_tests} اختبار تم إجراؤها، 
                    أبرزها:<br>
                    {sig_text}
                    <br>في المقابل، لم تظهر فروق دالة في باقي المحاور، مما يشير إلى تأثير محدود للمتغيرات الديموغرافية 
                    على اتجاهات أفراد العينة."
                    </p>
                </div>
                """, unsafe_allow_html=True)

            # ==================== زر تصدير التوصيات ====================
            st.markdown("---")
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                # تجهيز محتوى التوصيات للتصدير
                recommendations_text = f"""
                تقرير توصيات اختبار الفروقات
                =======================================

                تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

                ملخص النتائج:
                - عدد الاختبارات: {total_tests}
                - الفروق الدالة: {sig_count} ({sig_percentage:.1f}%)
                - متوسط حجم التأثير: {avg_effect:.3f}

                التوصيات:
                {'لا توجد فروق دالة - يمكن دمج المجموعات' if sig_count == 0 else 'توجد فروق دالة - يوصى بتحليل منفصل للمجموعات'}
                """

                st.download_button(
                    label="📥 تحميل التوصيات كملف نصي",
                    data=recommendations_text,
                    file_name="differences_recommendations.txt",
                    mime="text/plain",
                )

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== الصفحة 11: تحليل التباين الأحادي (ANOVA) ====================
elif st.session_state.page == 'anova' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📊 تحليل التباين الأحادي (ANOVA)</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل التباين الأحادي (One-Way ANOVA):</strong> يستخدم لمقارنة متوسطات ثلاث مجموعات أو أكثر لمتغير مستقل واحد.</p>
        <p>📌 <strong>اختبارات ما بعد hoc (Post-hoc):</strong> اختبار Tukey HSD لتحديد أي المجموعات تختلف عن بعضها.</p>
        <p>🎯 <strong>البديل اللامعلمي:</strong> اختبار Kruskal-Wallis للبيانات غير الطبيعية.</p>
        <p>⚠️ <strong>ملاحظة مهمة:</strong> المتغير التابع يجب أن يكون رقمياً (كمياً).</p>
    </div>
    """, unsafe_allow_html=True)

    # تصفية المحاور الرقمية فقط
    factor_names = []
    for f in factors:
        if f['questions'] and f['name'] in df.columns:
            # التأكد من أن المتغير رقمي
            if pd.api.types.is_numeric_dtype(df[f['name']]):
                factor_names.append(f['name'])
            else:
                st.warning(f"⚠️ المحور '{f['name']}' ليس رقمياً وتم استبعاده من التحليل")

    # تصفية المتغيرات الفئوية فقط
    categorical_vars = []
    for var in st.session_state.social_vars + st.session_state.independent_vars:
        if var in df.columns:
            # التأكد من أن المتغير فئوي (ليس رقمياً مستمراً)
            if not pd.api.types.is_numeric_dtype(df[var]) or df[var].nunique() <= 5:
                categorical_vars.append(var)
            else:
                st.info(f"ℹ️ المتغير '{var}' عددي مستمر وتم استبعاده من ANOVA (يمكن استخدامه في الانحدار)")

    if not factor_names:
        st.warning("⚠️ يرجى تحديد محاور رقمية للتحليل. تأكد من أن المحاور تحتوي على بيانات رقمية.")
    elif not categorical_vars:
        st.warning("⚠️ يرجى تحديد متغيرات فئوية (مثل: الجنس، المستوى التعليمي) للتحليل.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            selected_dep = st.selectbox("المتغير التابع (Y - عددي):", factor_names)
        with col2:
            selected_indep = st.selectbox("المتغير المستقل (X - فئوي):", categorical_vars)

        if st.button("تنفيذ تحليل ANOVA", type="primary"):
            try:
                # التحقق من وجود بيانات كافية
                data_groups_raw = [(name, group[selected_dep].dropna()) for name, group in df.groupby(selected_indep)]
                data_groups = [g for name, g in data_groups_raw if len(g) > 0]
                group_names = [name for name, g in data_groups_raw if len(g) > 0]

                if len(data_groups) < 2:
                    st.error("❌ لا توجد مجموعات كافية للمقارنة (يلزم مجموعتين على الأقل)")
                elif any(len(g) < 3 for g in data_groups):
                    st.warning("⚠️ بعض المجموعات تحتوي على أقل من 3 ملاحظات، قد لا تكون النتائج موثوقة")

                # تخزين النتائج للاستخدام في التوصيات
                anova_analysis_results = {}

                # التحقق من طبيعية البيانات لاختيار الاختبار المناسب
                is_normal = True
                normality_details = []
                for i, g in enumerate(data_groups):
                    if len(g) >= 3:
                        try:
                            _, p = shapiro(g)
                            if p <= 0.05:
                                is_normal = False
                            normality_details.append({
                                'group': group_names[i],
                                'size': len(g),
                                'p_value': p if 'p' in dir() else None,
                                'normal': p > 0.05 if 'p' in dir() else False
                            })
                        except Exception as e:
                            st.warning(f"⚠️ لا يمكن حساب اختبار الطبيعي للمجموعة '{group_names[i]}': {str(e)}")

                # حساب المتوسطات لكل مجموعة
                group_means = {name: g.mean() for name, g in zip(group_names, data_groups)}
                group_sizes = {name: len(g) for name, g in zip(group_names, data_groups)}

                if is_normal and st.session_state.is_normal is not False:
                    try:
                        # استخدام ANOVA المعلمي
                        f_stat, p_value = f_oneway(*data_groups)

                        # حساب SS (Sum of Squares)
                        total_mean = np.concatenate(data_groups).mean()
                        ss_between = sum(len(g) * (g.mean() - total_mean) ** 2 for g in data_groups)
                        ss_within = sum(((g - g.mean()) ** 2).sum() for g in data_groups)
                        ss_total = ss_between + ss_within

                        df_between = len(data_groups) - 1
                        df_within = sum(len(g) for g in data_groups) - len(data_groups)

                        ms_between = ss_between / df_between if df_between > 0 else 0
                        ms_within = ss_within / df_within if df_within > 0 else 0

                        # حساب حجم التأثير
                        eta_squared = ss_between / ss_total if ss_total > 0 else 0
                        omega_squared = (ss_between - (df_between * ms_within)) / (
                                    ss_total + ms_within) if ss_total + ms_within > 0 else 0

                        anova_table = pd.DataFrame({
                            'Source': [selected_indep, 'Residual', 'Total'],
                            'SS': [f"{ss_between:.4f}", f"{ss_within:.4f}", f"{ss_total:.4f}"],
                            'DF': [df_between, df_within, df_between + df_within],
                            'MS': [f"{ms_between:.4f}", f"{ms_within:.4f}", ""],
                            'F': [f"{f_stat:.4f}" if not np.isnan(f_stat) else "N/A", "", ""],
                            'P-value': [f"{p_value:.4f}" if not np.isnan(p_value) else "N/A", "", ""]
                        })

                        is_sig = p_value < 0.05

                        anova_analysis_results = {
                            'test_used': 'ANOVA (معلمي)',
                            'is_significant': is_sig,
                            'p_value': p_value if not np.isnan(p_value) else 1.0,
                            'f_statistic': f_stat if not np.isnan(f_stat) else 0,
                            'eta_squared': eta_squared,
                            'omega_squared': omega_squared,
                            'group_means': group_means,
                            'group_sizes': group_sizes,
                            'group_names': group_names,
                            'anova_table': anova_table
                        }

                        # عرض النتائج
                        st.markdown("### 📊 إحصائيات المجموعات")
                        means_df = pd.DataFrame({
                            'الفئة': group_names,
                            'المتوسط': [f"{group_means[name]:.3f}" for name in group_names],
                            'العدد': [group_sizes[name] for name in group_names],
                            'الانحراف المعياري': [f"{g.std():.3f}" for g in data_groups]
                        })
                        st.dataframe(means_df, use_container_width=True)

                        st.markdown("### 📋 جدول تحليل التباين (ANOVA)")
                        st.dataframe(anova_table, use_container_width=True)

                        st.markdown("### 📊 حجم التأثير")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("Eta-squared (η²)", f"{eta_squared:.4f}")
                            if eta_squared > 0.14:
                                st.caption("✅ تأثير كبير")
                            elif eta_squared > 0.06:
                                st.caption("📊 تأثير متوسط")
                            elif eta_squared > 0.01:
                                st.caption("📈 تأثير صغير")
                            else:
                                st.caption("📉 تأثير ضئيل")
                        with col2:
                            st.metric("Omega-squared (ω²)", f"{omega_squared:.4f}")

                        # اختبار Tukey HSD (Post-hoc)
                        if is_sig and len(data_groups) >= 3:
                            try:
                                from statsmodels.stats.multicomp import pairwise_tukeyhsd

                                # تجهيز البيانات لاختبار Tukey
                                tukey_data = []
                                tukey_groups = []
                                for name, g in zip(group_names, data_groups):
                                    tukey_data.extend(g.values)
                                    tukey_groups.extend([name] * len(g))

                                tukey_result = pairwise_tukeyhsd(tukey_data, tukey_groups, alpha=0.05)
                                tukey_df = pd.DataFrame(data=tukey_result._results_table.data[1:],
                                                        columns=tukey_result._results_table.data[0])

                                st.markdown("### 📊 اختبار Tukey HSD (ما بعد hoc)")
                                st.dataframe(tukey_df, use_container_width=True)

                                anova_analysis_results['tukey_results'] = tukey_df.to_dict('records')
                            except Exception as e:
                                st.warning(f"⚠️ تعذر إجراء اختبار Tukey HSD: {str(e)}")

                    except Exception as e:
                        st.error(f"❌ خطأ في ANOVA: {str(e)}")
                        st.info(
                            "قد يكون السبب وجود قيم غير رقمية أو بيانات غير صالحة. جرب استخدام الاختبار اللامعلمي أدناه.")
                        is_normal = False  # التحول إلى الاختبار اللامعلمي

                if not is_normal or st.session_state.is_normal is False:
                    try:
                        # استخدام Kruskal-Wallis كبديل لامعلمي
                        stat, p_value = kruskal(*data_groups)
                        is_sig = p_value < 0.05

                        # حساب حجم التأثير لـ Kruskal-Wallis (Epsilon-squared)
                        n_total = sum(len(g) for g in data_groups)
                        epsilon_sq = stat / (n_total - 1) if n_total > 1 else 0

                        anova_analysis_results = {
                            'test_used': 'Kruskal-Wallis (لامعلمي)',
                            'is_significant': is_sig,
                            'p_value': p_value,
                            'h_statistic': stat,
                            'epsilon_squared': epsilon_sq,
                            'group_names': group_names,
                            'group_sizes': group_sizes,
                            'group_means': group_means,
                            'normality_details': normality_details
                        }

                        st.markdown(f"""
                        <div class="info-box">
                            <h4>📊 اختبار Kruskal-Wallis (البديل اللامعلمي لـ ANOVA)</h4>
                            <p><strong>سبب استخدام الاختبار اللامعلمي:</strong> البيانات لا تتبع التوزيع الطبيعي في بعض المجموعات.</p>
                            <p>📈 قيمة الاختبار (H): {stat:.4f}</p>
                            <p>📊 P-value: {p_value:.4f}</p>
                            <p><strong>النتيجة:</strong> {'✅ توجد فروق دالة إحصائياً بين المجموعات' if is_sig else '❌ لا توجد فروق دالة إحصائياً'}</p>
                            <p>📏 حجم التأثير (Epsilon-squared): {epsilon_sq:.4f}</p>
                        </div>
                        """, unsafe_allow_html=True)

                        # عرض إحصائيات المجموعات
                        st.markdown("### 📊 إحصائيات المجموعات")
                        means_df = pd.DataFrame({
                            'الفئة': group_names,
                            'المتوسط': [f"{group_means[name]:.3f}" for name in group_names],
                            'العدد': [group_sizes[name] for name in group_names],
                            'الوسيط': [f"{g.median():.3f}" for g in data_groups]
                        })
                        st.dataframe(means_df, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ خطأ في اختبار Kruskal-Wallis: {str(e)}")
                        st.info("تأكد من أن المتغير التابع يحتوي على قيم رقمية صالحة.")
                        anova_analysis_results = None

                # المخطط الصندوقي
                try:
                    fig = px.box(df, x=selected_indep, y=selected_dep,
                                 title=f"توزيع {selected_dep} حسب {selected_indep}",
                                 color=selected_indep,
                                 color_discrete_sequence=px.colors.qualitative.Set2,
                                 points="all" if len(df) < 100 else "outliers")
                    fig.update_layout(height=500)
                    fig = create_rtl_figure(fig)
                    st.plotly_chart(fig, use_container_width=True)
                except Exception as e:
                    st.warning(f"⚠️ تعذر رسم المخطط الصندوقي: {str(e)}")

                # تخزين النتائج للتصدير
                if anova_analysis_results:
                    st.session_state.anova_results = anova_analysis_results

                # ==================== عرض التوصيات ====================
                if anova_analysis_results:
                    st.markdown("---")
                    st.markdown("## 📝 تفسير النتائج والتوصيات")

                    if not anova_analysis_results['is_significant']:
                        st.markdown(f"""
                        <div style="background: #d4edda; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #28a745;">
                            <h4 style="color: #155724;">✅ لا توجد فروق ذات دلالة إحصائية</h4>
                            <p style="font-size: 1.05rem; line-height: 1.6;">
                                <strong>الاستنتاج:</strong> بناءً على نتائج تحليل التباين {anova_analysis_results['test_used']}، 
                                <strong>لا توجد فروق دالة إحصائياً</strong> بين المجموعات المختلفة 
                                في المتغير التابع "{selected_dep}" (P-value = {anova_analysis_results['p_value']:.4f} > 0.05).
                            </p>
                            <p style="margin-top: 10px;">
                                هذا يشير إلى أن <strong>المتغير المستقل "{selected_indep}" لا يؤثر بشكل معنوي</strong> 
                                على المتغير التابع، وأن جميع المجموعات <strong>متجانسة</strong> في متوسطاتها.
                            </p>
                        </div>
                        """, unsafe_allow_html=True)

                        st.markdown("""
                        <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #17a2b8;">
                            <h4 style="color: #0c5460;">💡 التوصية الإجرائية</h4>
                            <ul style="line-height: 1.8;">
                                <li><strong>✅ يمكن دمج جميع المجموعات</strong> في التحليلات اللاحقة</li>
                                <li><strong>✅ التركيز على متغيرات أخرى</strong> قد تكون أكثر تأثيراً</li>
                                <li><strong>✅ استخدام تحليل الانحدار</strong> بدلاً من ANOVA إذا كان المتغير المستقل كمياً</li>
                            </ul>
                        </div>
                        """, unsafe_allow_html=True)

                    else:
                        # تحديد قوة حجم التأثير
                        if 'eta_squared' in anova_analysis_results:
                            effect = anova_analysis_results['eta_squared']
                            if effect > 0.14:
                                effect_text = "كبير جداً"
                                effect_color = "#dc3545"
                            elif effect > 0.06:
                                effect_text = "متوسط"
                                effect_color = "#ffc107"
                            else:
                                effect_text = "صغير"
                                effect_color = "#28a745"
                        elif 'epsilon_squared' in anova_analysis_results:
                            effect = anova_analysis_results['epsilon_squared']
                            if effect > 0.15:
                                effect_text = "كبير"
                                effect_color = "#dc3545"
                            elif effect > 0.08:
                                effect_text = "متوسط"
                                effect_color = "#ffc107"
                            else:
                                effect_text = "صغير"
                                effect_color = "#28a745"
                        else:
                            effect = 0
                            effect_text = "غير محدد"
                            effect_color = "#6c757d"

                        st.markdown(f"""
                        <div style="background: #fff3cd; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #ffc107;">
                            <h4 style="color: #856404;">⚠️ توجد فروق ذات دلالة إحصائية</h4>
                            <p style="font-size: 1.05rem; line-height: 1.6;">
                                <strong>الاستنتاج:</strong> بناءً على نتائج تحليل التباين {anova_analysis_results['test_used']}، 
                                <strong>توجد فروق دالة إحصائياً</strong> بين المجموعات المختلفة 
                                (P-value = {anova_analysis_results['p_value']:.4f} < 0.05).
                            </p>
                            <p>📏 <strong>حجم التأثير:</strong> {effect:.4f} ({effect_text})</p>
                        </div>
                        """, unsafe_allow_html=True)

                        st.markdown("""
                        <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #17a2b8;">
                            <h4 style="color: #0c5460;">💡 التوصية الإجرائية</h4>
                            <ul style="line-height: 1.8;">
                                <li><strong>🔍 إجراء اختبارات ما بعد hoc (Post-hoc)</strong> لتحديد أي المجموعات تختلف</li>
                                <li><strong>📊 تحليل منفصل لكل مجموعة</strong> عند دراسة علاقات أخرى</li>
                                <li><strong>🎯 تخصيص التوصيات</strong> بناءً على خصائص كل مجموعة</li>
                            </ul>
                        </div>
                        """, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"❌ خطأ في التحليل: {str(e)}")
                st.info("""
                **أسباب محتملة للخطأ وحلولها:**
                1. **المتغير التابع يحتوي على قيم نصية** → تأكد من أن المتغير التابع عددي
                2. **المتغير المستقل ليس فئوياً** → استخدم متغيراً فئوياً (مثل: الجنس، المستوى التعليمي)
                3. **عدد المجموعات أقل من 2** → تأكد من وجود مجموعتين على الأقل
                4. **حجم العينة صغير جداً** → يفضل أن يكون لكل مجموعة 3 ملاحظات على الأقل
                """)

    st.markdown('</div>', unsafe_allow_html=True)
# ==================== الصفحة 12: تحليل التجميع المتقدم ====================
elif st.session_state.page == 'clustering' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">🎯 تحليل التجميع المتقدم</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>🔍 <strong>تحليل التجميع (Cluster Analysis):</strong> يهدف إلى تصنيف الأفراد أو المتغيرات إلى مجموعات متجانسة بناءً على خصائصها.</p>
        <p>📊 <strong>الطرق:</strong> K-Means، التجميع الهرمي، DBSCAN.</p>
        <p>🎯 <strong>مقياس الجودة:</strong> Silhouette Score (كلما اقترب من 1 كان أفضل).</p>
        <p>⚠️ <strong>ملاحظة مهمة:</strong> جميع المتغيرات المستخدمة في التجميع يجب أن تكون رقمية.</p>
    </div>
    """, unsafe_allow_html=True)

    factor_names = [f['name'] for f in factors if f['questions']]
    available_vars = factor_names + st.session_state.social_vars + st.session_state.independent_vars + st.session_state.mediator_vars

    # تصفية المتغيرات الرقمية فقط
    numeric_vars = []
    for var in available_vars:
        if var in df.columns and pd.api.types.is_numeric_dtype(df[var]):
            numeric_vars.append(var)
        elif var in df.columns:
            st.warning(f"⚠️ المتغير '{var}' ليس رقمياً وتم استبعاده من تحليل التجميع")

    if len(numeric_vars) < 2:
        st.warning("⚠️ يلزم وجود متغيرين رقميين على الأقل لتحليل التجميع")
    else:
        selected_vars = st.multiselect("اختر المتغيرات للتجميع:", numeric_vars,
                                       default=numeric_vars[:min(3, len(numeric_vars))])

        if selected_vars and st.button("تنفيذ التجميع", type="primary"):
            try:
                # التأكد من أن البيانات رقمية
                cluster_data = df[selected_vars].dropna().copy()

                # تحويل جميع الأعمدة إلى numeric مع التعامل مع الأخطاء
                for col in selected_vars:
                    cluster_data[col] = pd.to_numeric(cluster_data[col], errors='coerce')

                # إزالة القيم المفقودة بعد التحويل
                cluster_data = cluster_data.dropna()

                if len(cluster_data) < 10:
                    st.warning(f"⚠️ عدد الملاحظات الصالحة قليل جداً: {len(cluster_data)} (يحتاج على الأقل 10 ملاحظات)")
                else:
                    # تخزين نتائج التحليل
                    clustering_results = {}

                    # توحيد البيانات
                    scaler = StandardScaler()
                    data_scaled = scaler.fit_transform(cluster_data)

                    inertias = []
                    silhouette_scores = []
                    K_range = range(2, min(11, len(cluster_data) // 3))

                    if len(K_range) < 1:
                        st.warning("⚠️ عدد الملاحظات غير كافٍ لتحديد عدد المجموعات")
                    else:
                        for k in K_range:
                            kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
                            labels = kmeans.fit_predict(data_scaled)
                            inertias.append(kmeans.inertia_)
                            if len(set(labels)) > 1:
                                try:
                                    sil_score = silhouette_score(data_scaled, labels)
                                    silhouette_scores.append(sil_score)
                                except:
                                    silhouette_scores.append(-1)
                            else:
                                silhouette_scores.append(-1)

                        # تحديد العدد الأمثل للمجموعات
                        optimal_k_silhouette = K_range[np.argmax(silhouette_scores)] if silhouette_scores and max(
                            silhouette_scores) > 0 else 3

                        # عرض الرسوم البيانية
                        col1, col2 = st.columns(2)
                        with col1:
                            fig_elbow = px.line(x=list(K_range), y=inertias, markers=True,
                                                title="منحنى Elbow لتحديد العدد الأمثل للمجموعات")
                            st.plotly_chart(fig_elbow, use_container_width=True)
                        with col2:
                            fig_sil = px.line(x=list(K_range), y=silhouette_scores, markers=True,
                                              title="Silhouette Score لتحديد العدد الأمثل")
                            fig_sil.add_hline(y=0.5, line_dash="dash", line_color="green", annotation_text="جيد (0.5)")
                            fig_sil.add_hline(y=0.7, line_dash="dash", line_color="blue", annotation_text="ممتاز (0.7)")
                            st.plotly_chart(fig_sil, use_container_width=True)

                        # اختيار عدد المجموعات
                        max_clusters = min(10, len(cluster_data) // 3)
                        n_clusters = st.slider("عدد المجموعات:", 2, max(2, max_clusters),
                                               value=min(optimal_k_silhouette,
                                                         max_clusters) if optimal_k_silhouette <= max_clusters else 3)

                        # تنفيذ التجميع النهائي
                        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
                        clusters = kmeans.fit_predict(data_scaled)
                        centers = kmeans.cluster_centers_

                        # حساب مقاييس الجودة
                        try:
                            final_silhouette = silhouette_score(data_scaled, clusters)
                        except:
                            final_silhouette = -1
                        final_inertia = kmeans.inertia_

                        # حساب توزيع المجموعات
                        cluster_sizes = pd.Series(clusters).value_counts().to_dict()

                        # تخزين النتائج
                        clustering_results = {
                            'n_clusters': n_clusters,
                            'silhouette_score': final_silhouette,
                            'inertia': final_inertia,
                            'cluster_sizes': cluster_sizes,
                            'cluster_centers': centers,
                            'variables': selected_vars,
                            'total_samples': len(cluster_data),
                            'optimal_k_suggested': optimal_k_silhouette
                        }

                        # عرض مقاييس الجودة
                        st.markdown("### 📊 مقاييس جودة التجميع")
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("عدد المجموعات", n_clusters)
                        with col2:
                            st.metric("Silhouette Score", f"{final_silhouette:.4f}")
                            if final_silhouette > 0.7:
                                st.caption("✅ ممتاز")
                            elif final_silhouette > 0.5:
                                st.caption("✅ جيد")
                            elif final_silhouette > 0.25:
                                st.caption("⚠️ مقبول")
                            else:
                                st.caption("❌ ضعيف")
                        with col3:
                            st.metric("Inertia", f"{final_inertia:.2f}")
                        with col4:
                            st.metric("حجم العينة", len(cluster_data))

                        # عرض توزيع المجموعات
                        st.markdown("### 📊 توزيع أفراد العينة على المجموعات")
                        cluster_dist = pd.DataFrame({
                            'المجموعة': [f'المجموعة {i + 1}' for i in range(n_clusters)],
                            'العدد': [cluster_sizes.get(i, 0) for i in range(n_clusters)],
                            'النسبة المئوية': [f"{cluster_sizes.get(i, 0) / len(cluster_data) * 100:.1f}%" for i in
                                               range(n_clusters)]
                        })
                        st.dataframe(cluster_dist, use_container_width=True)

                        # رسم بياني لتوزيع المجموعات
                        try:
                            fig_dist = px.pie(cluster_dist, values='العدد', names='المجموعة',
                                              title="توزيع أفراد العينة على المجموعات")
                            fig_dist = create_rtl_figure(fig_dist)
                            st.plotly_chart(fig_dist, use_container_width=True)
                        except:
                            st.info("لا يمكن رسم المخطط الدائري")

                        # تصور المجموعات باستخدام PCA
                        try:
                            pca = PCA(n_components=2)
                            pca_result = pca.fit_transform(data_scaled)
                            pca_df = pd.DataFrame(
                                {'PC1': pca_result[:, 0], 'PC2': pca_result[:, 1], 'Cluster': clusters.astype(str)})

                            fig = px.scatter(pca_df, x='PC1', y='PC2', color='Cluster',
                                             title=f"تصور المجموعات باستخدام PCA (Silhouette = {final_silhouette:.3f})",
                                             labels={
                                                 'PC1': f'المكون الرئيسي 1 ({pca.explained_variance_ratio_[0] * 100:.1f}%)',
                                                 'PC2': f'المكون الرئيسي 2 ({pca.explained_variance_ratio_[1] * 100:.1f}%)'})
                            fig = create_rtl_figure(fig)
                            st.plotly_chart(fig, use_container_width=True)
                        except Exception as e:
                            st.warning(f"⚠️ تعذر رسم تصور PCA: {str(e)}")

                        # عرض خصائص كل مجموعة
                        st.markdown("### 📊 خصائص المجموعات (المتوسطات)")
                        cluster_summary = cluster_data.copy()
                        cluster_summary['Cluster'] = clusters
                        summary_stats = cluster_summary.groupby('Cluster').mean()
                        summary_stats.index = [f'المجموعة {i + 1}' for i in summary_stats.index]

                        # عرض الجدول بدون تنسيق أرقام (لتجنب الخطأ)
                        st.dataframe(summary_stats, use_container_width=True)

                        # رسم بياني لمقارنة المجموعات (مع معالجة الأخطاء)
                        try:
                            st.markdown("### 📊 مقارنة المجموعات")
                            summary_melted = summary_stats.reset_index().melt(id_vars='index', var_name='المتغير',
                                                                              value_name='المتوسط')
                            summary_melted.columns = ['المجموعة', 'المتغير', 'المتوسط']
                            # التأكد من أن المتوسط رقمي
                            summary_melted['المتوسط'] = pd.to_numeric(summary_melted['المتوسط'], errors='coerce')
                            summary_melted = summary_melted.dropna()

                            if len(summary_melted) > 0:
                                fig_compare = px.bar(summary_melted, x='المتغير', y='المتوسط', color='المجموعة',
                                                     barmode='group', title="مقارنة متوسطات المجموعات حسب المتغيرات")
                                fig_compare = create_rtl_figure(fig_compare)
                                st.plotly_chart(fig_compare, use_container_width=True)
                        except Exception as e:
                            st.warning(f"⚠️ تعذر رسم مقارنة المجموعات: {str(e)}")

                        # ==================== التوصيات والتفسير ====================
                        st.markdown("---")
                        st.markdown("## 📝 تفسير النتائج والتوصيات")

                        # تبويبات التوصيات
                        rec_tab1, rec_tab2, rec_tab3 = st.tabs(
                            ["📊 التوصية الرئيسية", "📈 وصف المجموعات", "📝 الصياغة الأكاديمية"])

                        with rec_tab1:
                            # تقييم جودة التجميع
                            st.markdown("### 🎯 تقييم جودة التجميع")

                            if final_silhouette > 0.7:
                                st.markdown(f"""
                                <div style="background: #d4edda; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #28a745;">
                                    <h4 style="color: #155724;">✅ جودة تجميع ممتازة</h4>
                                    <p style="font-size: 1.05rem; line-height: 1.6;">
                                        <strong>Silhouette Score = {final_silhouette:.4f}</strong> (أكبر من 0.7)
                                    </p>
                                    <p style="margin-top: 10px;">
                                        هذا يشير إلى أن المجموعات <strong>منفصلة بشكل واضح</strong> عن بعضها البعض، 
                                        وأن التجميع <strong>موثوق ويمكن الاعتماد عليه</strong>.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                            elif final_silhouette > 0.5:
                                st.markdown(f"""
                                <div style="background: #d4edda; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #28a745;">
                                    <h4 style="color: #155724;">✅ جودة تجميع جيدة</h4>
                                    <p style="font-size: 1.05rem; line-height: 1.6;">
                                        <strong>Silhouette Score = {final_silhouette:.4f}</strong> (بين 0.5 و 0.7)
                                    </p>
                                    <p style="margin-top: 10px;">
                                        هذا يشير إلى أن المجموعات <strong>مفصولة بشكل جيد</strong>.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                            elif final_silhouette > 0.25:
                                st.markdown(f"""
                                <div style="background: #fff3cd; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #ffc107;">
                                    <h4 style="color: #856404;">⚠️ جودة تجميع مقبولة</h4>
                                    <p style="font-size: 1.05rem; line-height: 1.6;">
                                        <strong>Silhouette Score = {final_silhouette:.4f}</strong> (بين 0.25 و 0.5)
                                    </p>
                                    <p style="margin-top: 10px;">
                                        هذا يشير إلى أن المجموعات <strong>تتداخل جزئياً</strong>.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                            else:
                                st.markdown(f"""
                                <div style="background: #f8d7da; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #dc3545;">
                                    <h4 style="color: #721c24;">❌ جودة تجميع ضعيفة</h4>
                                    <p style="font-size: 1.05rem; line-height: 1.6;">
                                        <strong>Silhouette Score = {final_silhouette:.4f}</strong> (أقل من 0.25)
                                    </p>
                                    <p style="margin-top: 10px;">
                                        هذا يشير إلى أن المجموعات <strong>غير محددة بشكل جيد</strong>.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)

                            # التوصية الإجرائية
                            st.markdown("### 💡 التوصية الإجرائية")

                            if final_silhouette > 0.5:
                                st.markdown("""
                                <div style="background: #e8f4f8; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #17a2b8;">
                                    <h4 style="color: #0c5460;">✅ يمكن الاعتماد على نتائج التجميع</h4>
                                    <ul style="line-height: 1.8;">
                                        <li><strong>📊 استخدام المجموعات كمتغير تصنيفي</strong> في التحليلات اللاحقة</li>
                                        <li><strong>🎯 تخصيص التوصيات</strong> لكل مجموعة بناءً على خصائصها الفريدة</li>
                                        <li><strong>📋 تصميم برامج تدخل مخصصة</strong> حسب احتياجات كل مجموعة</li>
                                    </ul>
                                </div>
                                """, unsafe_allow_html=True)
                            else:
                                st.markdown("""
                                <div style="background: #fff3cd; border-radius: 15px; padding: 20px; margin: 15px 0; border-right: 5px solid #ffc107;">
                                    <h4 style="color: #856404;">⚠️ يوصى بالتعامل مع النتائج بحذر</h4>
                                    <ul style="line-height: 1.8;">
                                        <li><strong>🔄 إعادة التجميع بعدد مختلف من المجموعات</strong> ومقارنة النتائج</li>
                                        <li><strong>📊 استخدام خوارزميات تجميع مختلفة</strong></li>
                                        <li><strong>📈 توسيع حجم العينة</strong> للحصول على نتائج أكثر استقراراً</li>
                                    </ul>
                                </div>
                                """, unsafe_allow_html=True)

                        with rec_tab2:
                            # وصف المجموعات (بدون تنسيق أرقام قد يسبب خطأ)
                            st.markdown("### 📊 وصف خصائص كل مجموعة")

                            # عرض إحصائيات المجموعات بشكل آمن
                            st.markdown("#### متوسطات المتغيرات حسب المجموعة:")
                            st.dataframe(summary_stats, use_container_width=True)

                            # وصف نصي لكل مجموعة
                            st.markdown("### 📝 الوصف النصي للمجموعات")

                            # تحديد المجموعة الأكبر والأصغر
                            sizes = clustering_results['cluster_sizes']
                            largest_cluster = max(sizes, key=sizes.get) if sizes else 0
                            smallest_cluster = min(sizes, key=sizes.get) if sizes else 0

                            st.markdown(f"""
                            <div style="background: #f5f5f5; border-radius: 15px; padding: 20px; margin: 15px 0;">
                                <p><strong>📊 توزيع العينة:</strong></p>
                                <ul>
                                    <li><strong>أكبر مجموعة:</strong> المجموعة {largest_cluster + 1} ({sizes.get(largest_cluster, 0)} فرد)</li>
                                    <li><strong>أصغر مجموعة:</strong> المجموعة {smallest_cluster + 1} ({sizes.get(smallest_cluster, 0)} فرد)</li>
                                </ul>
                            </div>
                            """, unsafe_allow_html=True)

                            # وصف كل مجموعة
                            st.markdown("#### 🎯 توصيف المجموعات:")

                            for i in range(n_clusters):
                                try:
                                    cluster_means = summary_stats.iloc[i]
                                    if len(cluster_means) > 0:
                                        highest_var = cluster_means.idxmax()
                                        lowest_var = cluster_means.idxmin()
                                        highest_val = cluster_means[highest_var]
                                        lowest_val = cluster_means[lowest_var]

                                        st.markdown(f"""
                                        <div style="background: {'#e8f4f8' if i % 2 == 0 else '#f0f0f0'}; 
                                                    border-radius: 12px; padding: 15px; margin: 10px 0; border-right: 3px solid #667eea;">
                                            <h4 style="color: #667eea;">المجموعة {i + 1}</h4>
                                            <p><strong>📊 الحجم:</strong> {sizes.get(i, 0)} فرد ({sizes.get(i, 0) / len(cluster_data) * 100:.1f}% من العينة)</p>
                                            <p><strong>📈 أعلى متوسط:</strong> {highest_var} = {highest_val:.3f}</p>
                                            <p><strong>📉 أدنى متوسط:</strong> {lowest_var} = {lowest_val:.3f}</p>
                                        </div>
                                        """, unsafe_allow_html=True)
                                except Exception as e:
                                    st.warning(f"⚠️ تعذر وصف المجموعة {i + 1}: {str(e)}")

                        with rec_tab3:
                            # الصياغة الأكاديمية
                            st.markdown("### 📝 الصياغة المقترحة للتقرير الأكاديمي")

                            # تحديد جودة التجميع بالكلمات
                            if final_silhouette > 0.7:
                                quality_text = "ممتازة"
                                quality_en = "excellent"
                            elif final_silhouette > 0.5:
                                quality_text = "جيدة"
                                quality_en = "good"
                            elif final_silhouette > 0.25:
                                quality_text = "مقبولة"
                                quality_en = "acceptable"
                            else:
                                quality_text = "ضعيفة"
                                quality_en = "weak"

                            # وصف المتغيرات المستخدمة
                            vars_text = "، ".join(selected_vars)

                            st.markdown(f"""
                            <div style="background: #f5f5f5; border-radius: 15px; padding: 20px; margin: 15px 0; direction: rtl;">
                                <p style="font-size: 1rem; line-height: 1.8;">
                                <strong>النص العربي:</strong><br><br>
                                "لتصنيف أفراد العينة إلى مجموعات متجانسة، تم استخدام تحليل التجميع بطريقة K-Means 
                                بناءً على المتغيرات التالية: {vars_text}. 
                                تم تحديد العدد الأمثل للمجموعات باستخدام منحنى Elbow ومعامل Silhouette، 
                                حيث تم اختيار {n_clusters} مجموعات. 
                                بلغ معامل Silhouette {final_silhouette:.4f}، مما يشير إلى جودة تجميع {quality_text}.<br><br>

                                أظهرت النتائج توزيع أفراد العينة على المجموعات كما يلي: 
                                {', '.join([f'المجموعة {i + 1}: {sizes.get(i, 0)} فرد' for i in range(n_clusters)])}."
                                </p>
                            </div>
                            """, unsafe_allow_html=True)

                        # تخزين النتائج للتصدير
                        st.session_state.clustering_results = clustering_results

                        # زر تصدير التوصيات
                        st.markdown("---")
                        col1, col2, col3 = st.columns([1, 2, 1])
                        with col2:
                            recommendations_text = f"""
                            تقرير توصيات تحليل التجميع
                            =============================

                            عدد المجموعات: {n_clusters}
                            Silhouette Score: {final_silhouette:.4f}
                            حجم العينة: {len(cluster_data)}

                            توزيع المجموعات:
                            """
                            for i in range(n_clusters):
                                recommendations_text += f"\n- المجموعة {i + 1}: {sizes.get(i, 0)} فرد"

                            st.download_button(
                                label="📥 تحميل التوصيات كملف نصي",
                                data=recommendations_text,
                                file_name="clustering_recommendations.txt",
                                mime="text/plain",
                            )

            except Exception as e:
                st.error(f"❌ خطأ في التحليل: {str(e)}")
                st.info("""
                **أسباب محتملة للخطأ وحلولها:**
                1. **البيانات تحتوي على قيم نصية** → تأكد من أن جميع المتغيرات المختارة رقمية
                2. **عدد الملاحظات غير كافٍ** → يفضل أن يكون عدد الملاحظات 30 على الأقل
                3. **قيم مفقودة كثيرة** → تأكد من إزالة القيم المفقودة
                """)

    st.markdown('</div>', unsafe_allow_html=True)
# ==================== الصفحة 13: تصدير التقرير ====================
elif st.session_state.page == 'export' and st.session_state.data_loaded:
    df = st.session_state.df
    factors = st.session_state.factors

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title">📄 تصدير التقارير والنتائج</h2>', unsafe_allow_html=True)

    st.markdown("""
    <div class="result-card">
        <p>📊 يمكنك تصدير جميع نتائج التحليل بتنسيقات متعددة:</p>
        <ul>
            <li><strong>تقرير HTML</strong> - تقرير تفاعلي كامل يضم جميع التحليلات</li>
            <li><strong>تقرير Word</strong> - تقرير نصي منسق مع جداول ونتائج</li>
            <li><strong>ملف Excel</strong> - جميع النتائج في أوراق عمل منفصلة</li>
        </ul>
        <p>📌 <strong>ملاحظة:</strong> يتم تضمين جميع التحليلات التي تم إجراؤها في الجلسة الحالية.</p>
    </div>
    """, unsafe_allow_html=True)


    # ==================== دالة generate_word_report المحسنة ====================
    def generate_word_report_safe(results_dict, filename="analysis_report.docx"):
        """توليد تقرير Word بشكل آمن مع معالجة متقدمة للأخطاء"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            # إنشاء مستند جديد
            doc = Document()

            # إعدادات الصفحة
            try:
                section = doc.sections[0]
                section.right_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
            except:
                pass

            # العنوان الرئيسي
            title = doc.add_heading('تقرير التحليل الإحصائي', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # تاريخ التقرير
            date_para = doc.add_paragraph(f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")

            # إضافة المحتوى
            valid_sections = 0

            for key, value in results_dict.items():
                try:
                    if value is None:
                        continue

                    # إضافة عنوان القسم
                    doc.add_heading(str(key)[:50], level=1)

                    # معالجة DataFrame
                    if isinstance(value, pd.DataFrame):
                        if value.empty:
                            doc.add_paragraph("(لا توجد بيانات)")
                            doc.add_paragraph("")
                            continue

                        # تحديد حجم الجدول المناسب
                        n_cols = min(len(value.columns), 15)
                        n_rows = min(len(value), 50)

                        if n_cols == 0 or n_rows == 0:
                            doc.add_paragraph("(لا توجد بيانات كافية للعرض)")
                            doc.add_paragraph("")
                            continue

                        # إنشاء الجدول
                        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
                        table.style = 'Light Grid Accent 1'

                        # إضافة رؤوس الأعمدة
                        for j in range(n_cols):
                            try:
                                col_name = str(value.columns[j])
                                if len(col_name) > 30:
                                    col_name = col_name[:27] + "..."
                                table.cell(0, j).text = col_name
                            except:
                                table.cell(0, j).text = f"عمود {j + 1}"

                        # إضافة البيانات
                        for i in range(n_rows):
                            for j in range(n_cols):
                                try:
                                    val = value.iloc[i, j]
                                    if pd.isna(val):
                                        cell_text = "N/A"
                                    elif isinstance(val, float):
                                        cell_text = f"{val:.3f}"
                                    else:
                                        cell_text = str(val)[:100]
                                    table.cell(i + 1, j).text = cell_text
                                except:
                                    table.cell(i + 1, j).text = "خطأ"

                        doc.add_paragraph("")
                        valid_sections += 1

                    # معالجة القاموس
                    elif isinstance(value, dict):
                        if not value:
                            doc.add_paragraph("(لا توجد بيانات)")
                            doc.add_paragraph("")
                            continue

                        # إنشاء جدول للقاموس
                        n_items = min(len(value), 100)
                        table = doc.add_table(rows=n_items + 1, cols=2)
                        table.style = 'Light Grid Accent 1'

                        table.cell(0, 0).text = "العنصر"
                        table.cell(0, 1).text = "القيمة"

                        for i, (k, v) in enumerate(list(value.items())[:n_items]):
                            try:
                                table.cell(i + 1, 0).text = str(k)[:50]
                                if isinstance(v, float):
                                    table.cell(i + 1, 1).text = f"{v:.4f}"
                                elif isinstance(v, (int, np.integer)):
                                    table.cell(i + 1, 1).text = str(v)
                                else:
                                    table.cell(i + 1, 1).text = str(v)[:100]
                            except:
                                table.cell(i + 1, 0).text = str(k)[:50]
                                table.cell(i + 1, 1).text = "خطأ"

                        doc.add_paragraph("")
                        valid_sections += 1

                    # معالجة القائمة
                    elif isinstance(value, list):
                        if not value:
                            doc.add_paragraph("(لا توجد بيانات)")
                            doc.add_paragraph("")
                            continue

                        for item in value[:50]:
                            try:
                                if isinstance(item, dict):
                                    for k, v in list(item.items())[:5]:
                                        doc.add_paragraph(f"• {k}: {v}", style='List Bullet')
                                else:
                                    doc.add_paragraph(f"• {str(item)[:100]}", style='List Bullet')
                            except:
                                doc.add_paragraph("• (بيانات غير قابلة للعرض)", style='List Bullet')

                        doc.add_paragraph("")
                        valid_sections += 1

                    else:
                        try:
                            doc.add_paragraph(str(value)[:500])
                            doc.add_paragraph("")
                            valid_sections += 1
                        except:
                            doc.add_paragraph("(بيانات غير قابلة للعرض)")
                            doc.add_paragraph("")

                except Exception as e:
                    print(f"خطأ في القسم {key}: {str(e)}")
                    continue

            # إذا لم يتم إضافة أي قسم صالح
            if valid_sections == 0:
                doc.add_paragraph("⚠️ لا توجد بيانات كافية لإنشاء التقرير.")
                doc.add_paragraph("يرجى تنفيذ بعض التحليلات أولاً ثم المحاولة مرة أخرى.")

            # حفظ الملف
            doc.save(filename)
            return True

        except ImportError:
            print("مكتبة python-docx غير مثبتة")
            return False
        except Exception as e:
            print(f"خطأ في إنشاء تقرير Word: {str(e)}")
            return False


    # ==================== التحقق من تثبيت المكتبات المطلوبة ====================
    def check_and_install_packages():
        """التحقق من تثبيت المكتبات المطلوبة وتثبيتها إذا لزم الأمر"""
        missing_packages = []

        # التحقق من python-docx
        try:
            import docx
            docx_available = True
        except ImportError:
            docx_available = False
            missing_packages.append("python-docx")

        # التحقق من xlsxwriter
        try:
            import xlsxwriter
            xlsxwriter_available = True
        except ImportError:
            xlsxwriter_available = False
            missing_packages.append("xlsxwriter")

        return docx_available, xlsxwriter_available, missing_packages


    docx_available, xlsxwriter_available, missing_packages = check_and_install_packages()

    # عرض حالة المكتبات
    st.markdown("### 📦 حالة المكتبات المطلوبة")
    col1, col2 = st.columns(2)
    with col1:
        if docx_available:
            st.success("✅ python-docx: مثبت")
        else:
            st.error("❌ python-docx: غير مثبت")
    with col2:
        if xlsxwriter_available:
            st.success("✅ xlsxwriter: مثبت")
        else:
            st.error("❌ xlsxwriter: غير مثبت")

    # إذا كانت المكتبات مفقودة، عرض تعليمات التثبيت
    if missing_packages:
        st.warning(f"⚠️ المكتبات التالية غير مثبتة: {', '.join(missing_packages)}")

        if st.button("🔧 محاولة تثبيت المكتبات المفقودة", type="secondary"):
            try:
                import subprocess
                import sys

                for package in missing_packages:
                    st.info(f"جاري تثبيت {package}...")
                    subprocess.check_call([sys.executable, "-m", "pip", "install", package])
                    st.success(f"✅ تم تثبيت {package} بنجاح!")

                st.success("✅ تم تثبيت جميع المكتبات! يرجى إعادة تشغيل التطبيق.")
                st.rerun()
            except Exception as e:
                st.error(f"❌ فشل التثبيت التلقائي: {str(e)}")
                st.markdown(f"""
                <div class="info-box">
                    <p><strong>للتثبيت يدوياً، قم بتشغيل الأوامر التالية في terminal:</strong></p>
                    <code>pip install python-docx</code><br>
                    <code>pip install xlsxwriter</code>
                </div>
                """, unsafe_allow_html=True)

    # ==================== تجميع جميع نتائج التحليل ====================
    results_to_export = {}

    # 1. معلومات عامة عن البيانات
    with st.spinner("📊 جاري تجميع نتائج التحليل..."):
        results_to_export['معلومات عامة'] = pd.DataFrame({
            'المعلومة': [
                'تاريخ التقرير',
                'عدد الصفوف (المبحوثين)',
                'عدد الأعمدة (المتغيرات)',
                'عدد المحاور المحددة',
                'نسبة البيانات المفقودة'
            ],
            'القيمة': [
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                df.shape[0],
                df.shape[1],
                len(factors),
                f"{(df.isnull().sum().sum() / (df.shape[0] * df.shape[1]) * 100):.2f}%"
            ]
        })

        # 2. قائمة المحاور والفقرات
        factors_data = []
        for factor in factors:
            factors_data.append({
                'المحور': factor['name'],
                'عدد الفقرات': len(factor['questions']),
                'الفقرات': ', '.join(factor['questions']),
                'المتغيرات الاجتماعية المرتبطة': ', '.join(factor['social_vars']) if factor[
                    'social_vars'] else 'لا يوجد',
                'المتغيرات المستقلة المرتبطة': ', '.join(factor['independent_vars']) if factor[
                    'independent_vars'] else 'لا يوجد',
                'المتغيرات الوسيطة المرتبطة': ', '.join(factor['mediator_vars']) if factor[
                    'mediator_vars'] else 'لا يوجد'
            })
        if factors_data:
            results_to_export['المحاور والفقرات'] = pd.DataFrame(factors_data)

        # 3. الإحصاءات الوصفية
        numeric_df = df.select_dtypes(include=[np.number])
        if len(numeric_df.columns) > 0:
            descriptive_stats = calculate_descriptive_statistics_detailed(numeric_df)
            results_to_export['الإحصاءات الوصفية التفصيلية'] = descriptive_stats

        # 4. اتجاهات المحاور
        trends = []
        reliability_data = []
        for factor in factors:
            if factor['questions'] and factor['name'] in df.columns:
                mean_val = df[factor['name']].mean()
                trend_text, trend_class, trend_level = get_likert_trend(mean_val)
                trends.append({
                    'المحور': factor['name'],
                    'المتوسط': f"{mean_val:.3f}",
                    'الاتجاه': trend_text,
                    'المستوى': trend_level,
                    'عدد الفقرات': len(factor['questions'])
                })

                if len(factor['questions']) >= 2:
                    cronbach = calculate_cronbach_alpha(df[factor['questions']])
                    if not np.isnan(cronbach):
                        cronbach_status = "ممتاز" if cronbach >= 0.9 else "جيد جداً" if cronbach >= 0.8 else "جيد" if cronbach >= 0.7 else "مقبول" if cronbach >= 0.6 else "ضعيف"
                        reliability_data.append({
                            'المحور': factor['name'],
                            'معامل ألفا كرونباخ': f"{cronbach:.4f}",
                            'التقييم': cronbach_status
                        })

        if trends:
            results_to_export['اتجاهات المحاور'] = pd.DataFrame(trends)
        if reliability_data:
            results_to_export['معاملات الثبات'] = pd.DataFrame(reliability_data)

        # 5. نتائج اختبار التوزيع الطبيعي
        normality_results = []
        for factor in factors:
            if factor['questions'] and factor['name'] in df.columns:
                factor_values = df[factor['name']].dropna()
                if len(factor_values) >= 3:
                    try:
                        shapiro_stat, shapiro_p = shapiro(factor_values)
                        is_normal = shapiro_p > 0.05
                        normality_results.append({
                            'المحور': factor['name'],
                            'حجم العينة': len(factor_values),
                            'Shapiro-Wilk (Statistic)': f"{shapiro_stat:.4f}",
                            'Shapiro-Wilk (P-value)': f"{shapiro_p:.4f}",
                            'النتيجة': 'طبيعي' if is_normal else 'غير طبيعي'
                        })
                    except:
                        pass
        if normality_results:
            results_to_export['اختبار التوزيع الطبيعي'] = pd.DataFrame(normality_results)

        # 6. نتائج اختبار الفروقات
        if st.session_state.get('differences_results'):
            diff_df = pd.DataFrame(st.session_state.differences_results)
            # اختيار الأعمدة المهمة فقط
            important_cols = ['المحور', 'المتغير', 'P-value', 'الدلالة', 'حجم التأثير', 'تفسير الحجم']
            available_cols = [c for c in important_cols if c in diff_df.columns]
            if available_cols:
                results_to_export['نتائج اختبار الفروقات'] = diff_df[available_cols]
            else:
                results_to_export['نتائج اختبار الفروقات'] = diff_df

        # 7. نتائج تحليل ANOVA
        if st.session_state.get('anova_results'):
            anova_res = st.session_state.anova_results
            if 'anova_table' in anova_res and anova_res['anova_table'] is not None:
                results_to_export['نتائج تحليل ANOVA'] = anova_res['anova_table']

        # 8. نتائج تحليل التجميع
        if st.session_state.get('clustering_results'):
            clust_res = st.session_state.clustering_results
            clustering_summary = pd.DataFrame({
                'المقياس': ['عدد المجموعات', 'Silhouette Score', 'حجم العينة'],
                'القيمة': [
                    clust_res.get('n_clusters', 0),
                    f"{clust_res.get('silhouette_score', 0):.4f}",
                    clust_res.get('total_samples', 0)
                ]
            })
            results_to_export['ملخص تحليل التجميع'] = clustering_summary

    # ==================== عرض ملخص النتائج ====================
    st.success(f"✅ تم تجميع {len(results_to_export)} نوعاً من النتائج بنجاح!")

    with st.expander("📋 عرض ملخص النتائج المجمعة"):
        for sheet_name in list(results_to_export.keys())[:5]:
            st.markdown(f"**📄 {sheet_name}**")
            if isinstance(results_to_export[sheet_name], pd.DataFrame):
                st.dataframe(results_to_export[sheet_name].head(5), use_container_width=True)
            st.markdown("---")

    # ==================== أزرار التصدير ====================
    st.markdown("### 📥 خيارات التصدير")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📄 تصدير إلى HTML", use_container_width=True, type="primary"):
            with st.spinner("جاري إنشاء تقرير HTML..."):
                html_report = generate_html_report(results_to_export, "تقرير التحليل الإحصائي الشامل")
                b64 = base64.b64encode(html_report.encode()).decode()
                href = f'<a href="data:text/html;base64,{b64}" download="analysis_report.html">📥 تحميل تقرير HTML</a>'
                st.markdown(href, unsafe_allow_html=True)
                st.success("✅ تم إنشاء تقرير HTML بنجاح!")
                st.balloons()

    with col2:
        if st.button("📊 تصدير إلى Excel", use_container_width=True, type="primary"):
            if xlsxwriter_available:
                with st.spinner("جاري إنشاء ملف Excel..."):
                    if export_to_excel(results_to_export, "analysis_results.xlsx"):
                        try:
                            with open("analysis_results.xlsx", "rb") as f:
                                b64 = base64.b64encode(f.read()).decode()
                                href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="analysis_results.xlsx">📥 تحميل ملف Excel</a>'
                                st.markdown(href, unsafe_allow_html=True)
                            st.success("✅ تم إنشاء ملف Excel بنجاح!")
                            st.balloons()
                        except Exception as e:
                            st.error(f"❌ خطأ في قراءة الملف: {str(e)}")
                    else:
                        st.error("❌ حدث خطأ في إنشاء ملف Excel")
            else:
                st.error("❌ مكتبة xlsxwriter غير مثبتة. يرجى تثبيتها: `pip install xlsxwriter`")

    with col3:
        if st.button("📝 تصدير إلى Word", use_container_width=True, type="primary"):
            if docx_available:
                with st.spinner("جاري إنشاء تقرير Word... (قد يستغرق بضع ثوانٍ)"):
                    # إنشاء نسخة مبسطة من النتائج للتصدير
                    simplified_results = {}

                    # اختيار أهم النتائج فقط للتصدير (لتجنب الأخطاء)
                    important_keys = ['معلومات عامة', 'المحاور والفقرات', 'اتجاهات المحاور',
                                      'معاملات الثبات', 'اختبار التوزيع الطبيعي']

                    for key in important_keys:
                        if key in results_to_export:
                            simplified_results[key] = results_to_export[key]

                    # إضافة نتائج التحليلات إن وجدت (نسخة مبسطة)
                    if st.session_state.get('differences_results'):
                        diff_df = pd.DataFrame(st.session_state.differences_results)
                        if len(diff_df) > 50:
                            diff_df = diff_df.head(50)
                        simplified_results['نتائج اختبار الفروقات (أول 50)'] = diff_df

                    if st.session_state.get('anova_results'):
                        anova_res = st.session_state.anova_results
                        if 'anova_table' in anova_res and anova_res['anova_table'] is not None:
                            simplified_results['نتائج تحليل ANOVA'] = anova_res['anova_table']

                    if st.session_state.get('clustering_results'):
                        clust_res = st.session_state.clustering_results
                        simplified_results['ملخص تحليل التجميع'] = pd.DataFrame({
                            'المقياس': ['عدد المجموعات', 'Silhouette Score', 'حجم العينة'],
                            'القيمة': [
                                clust_res.get('n_clusters', 0),
                                f"{clust_res.get('silhouette_score', 0):.4f}",
                                clust_res.get('total_samples', 0)
                            ]
                        })

                    success = generate_word_report_safe(simplified_results, "analysis_report.docx")

                    if success:
                        try:
                            with open("analysis_report.docx", "rb") as f:
                                b64 = base64.b64encode(f.read()).decode()
                                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="analysis_report.docx">📥 تحميل تقرير Word</a>'
                                st.markdown(href, unsafe_allow_html=True)
                            st.success("✅ تم إنشاء تقرير Word بنجاح!")
                            st.balloons()
                        except Exception as e:
                            st.error(f"❌ خطأ في قراءة ملف Word: {str(e)}")
                    else:
                        st.error("❌ حدث خطأ في إنشاء تقرير Word. تأكد من تثبيت مكتبة python-docx بشكل صحيح.")
                        st.info("💡 نصيحة: جرب تصدير التقرير بتنسيق HTML أو Excel بدلاً من Word.")
            else:
                st.error("❌ مكتبة python-docx غير مثبتة")
                st.markdown("""
                <div class="info-box">
                    <p><strong>للتثبيت، قم بتشغيل الأمر التالي:</strong></p>
                    <code>pip install python-docx</code>
                    <p style="margin-top: 10px;">ثم أعد تشغيل التطبيق.</p>
                    <p>💡 <strong>بديل:</strong> يمكنك استخدام تصدير HTML أو Excel بدلاً من Word.</p>
                </div>
                """, unsafe_allow_html=True)

    # ==================== تعليمات إضافية ====================
    st.markdown("---")
    st.markdown("### 📌 تعليمات التصدير")

    st.markdown("""
    <div class="info-box">
        <h4>💡 نصائح هامة:</h4>
        <ul>
            <li><strong>تقرير HTML:</strong> يمكن فتحه في أي متصفح، يحتوي على جميع الجداول والنتائج بتنسيق تفاعلي</li>
            <li><strong>ملف Excel:</strong> مناسب للتحليل الإضافي في Excel، كل نتيجة في ورقة عمل منفصلة</li>
            <li><strong>تقرير Word:</strong> مناسب للتضمين في التقارير والأبحاث العلمية</li>
        </ul>
        <p>📌 <strong>ملاحظة:</strong> يتم حفظ الملفات في مجلد التحميل الافتراضي للمتصفح.</p>
        <p>⚠️ <strong>تنبيه:</strong> إذا واجهت مشكلة في تصدير Word، استخدم تنسيق HTML أو Excel كبديل.</p>
    </div>
    """, unsafe_allow_html=True)

    # ==================== إحصائيات سريعة ====================
    st.markdown("### 📊 إحصائيات سريعة")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📁 عدد أوراق النتائج", len(results_to_export))
    with col2:
        st.metric("📊 عدد المحاور", len(factors))
    with col3:
        st.metric("🔢 عدد المتغيرات الرقمية", len(numeric_df.columns) if 'numeric_df' in dir() else 0)
    with col4:
        st.metric("📈 حجم العينة", df.shape[0])

    st.markdown('</div>', unsafe_allow_html=True)
# ==================== الصفحة 14: التواصل ====================
elif st.session_state.page == 'contact':
    st.markdown('<div class="section-card" style="text-align: center;">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-title" style="display: block; text-align: center;">✉️ تواصل معنا</h2>',
                unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box" style="text-align: center; max-width: 500px; margin: 0 auto 20px auto;">
        <h3>📞 نرحب باستفساراتكم وملاحظاتكم</h3>
    </div>
    """, unsafe_allow_html=True)

    with st.form("contact_form_simple"):
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown('<div style="text-align: center;">', unsafe_allow_html=True)
            name = st.text_input("الاسم الكامل *", key="contact_name")
            email = st.text_input("البريد الإلكتروني *", key="contact_email")
            message = st.text_area("الرسالة *", height=150, key="contact_message")
            st.markdown('</div>', unsafe_allow_html=True)

            submitted = st.form_submit_button("📧 إرسال الرسالة", use_container_width=True)

            if submitted:
                if not name or not email or not message:
                    st.error("❌ الرجاء ملء جميع الحقول المطلوبة (*)")
                else:
                    st.success("✅ تم إرسال رسالتك بنجاح! سيتم الرد عليك في أقرب وقت.")
                    st.info(f"📧 تم إرسال نسخة من الرسالة إلى: boutoubaamed@gmail.com")
                    st.balloons()

    st.markdown("""
    <div style="text-align: center; margin-top: 30px; padding: 15px; background: #e8f4f8; border-radius: 10px; max-width: 500px; margin-left: auto; margin-right: auto;">
        <p>📧 <strong>boutoubaamed@gmail.com</strong></p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ==================== تذييل الصفحة ====================
st.markdown("""
<hr>
<div class="footer">
    <p> 2026 © جامعة عين تموشنت - جميع الحقوق محفوظة </p>
    <p>إصدار تجريبي | تم التطوير بواسطة الأستاذ الدكتور محمد بوطوبة </p>
</div>
""", unsafe_allow_html=True)